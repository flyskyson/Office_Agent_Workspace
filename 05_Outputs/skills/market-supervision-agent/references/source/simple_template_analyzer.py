#!/usr/bin/env python3
"""
Word模板分析工具 - 简化版
用于分析个体工商户登记申请书的Word模板结构
"""

import sys
from pathlib import Path
from typing import Dict, List, Tuple, Set
import json

try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table, _Cell
except ImportError:
    print("需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)

class SimpleTemplateAnalyzer:
    """Word模板分析器 - 简化版"""

    def __init__(self):
        self.placeholder_patterns = [
            "{{", "}}",  # 双花括号
            "[", "]",    # 方括号
            "<<", ">>",  # 双尖括号
            "__", "__",  # 双下划线
            "##", "##",  # 双井号
        ]

    def analyze_template(self, template_path: str) -> Dict:
        """
        分析Word模板结构

        Args:
            template_path: Word模板文件路径

        Returns:
            包含模板分析结果的字典
        """
        template_path = Path(template_path)
        if not template_path.exists():
            print(f"模板文件不存在: {template_path}")
            return {}

        print(f"\n正在分析模板: {template_path.name}")
        print("=" * 60)

        try:
            doc = Document(str(template_path))
            analysis = {
                "template_name": template_path.name,
                "template_size_kb": template_path.stat().st_size / 1024,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "sections_count": len(doc.sections),
                "placeholders_found": [],
                "suggested_fields": [],
                "template_structure": [],
                "extracted_text": []
            }

            # 分析段落
            print(f"\n段落分析:")
            print(f"  段落总数: {len(doc.paragraphs)}")

            placeholder_count = 0
            for i, paragraph in enumerate(doc.paragraphs[:20]):  # 只分析前20个段落
                if paragraph.text.strip():
                    analysis["extracted_text"].append({
                        "type": "paragraph",
                        "index": i,
                        "text": paragraph.text.strip(),
                        "style": paragraph.style.name if paragraph.style else "Normal"
                    })

                    # 检查占位符
                    placeholders = self._find_placeholders(paragraph.text)
                    if placeholders:
                        analysis["placeholders_found"].extend(placeholders)
                        placeholder_count += len(placeholders)
                        if placeholders:
                            print(f"  段落 {i+1}: 找到占位符 {placeholders}")

            # 分析表格
            print(f"\n表格分析:")
            print(f"  表格总数: {len(doc.tables)}")

            for i, table in enumerate(doc.tables):
                table_info = {
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "cells": []
                }

                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            table_info["cells"].append({
                                "row": row_idx,
                                "col": col_idx,
                                "text": cell_text
                            })

                            # 检查占位符
                            placeholders = self._find_placeholders(cell_text)
                            if placeholders:
                                analysis["placeholders_found"].extend(placeholders)
                                placeholder_count += len(placeholders)
                                if placeholders:
                                    print(f"  表格 {i+1}, 单元格({row_idx+1},{col_idx+1}): 找到占位符 {placeholders}")

                analysis["template_structure"].append({
                    "type": "table",
                    "info": table_info
                })

            # 去重占位符
            analysis["placeholders_found"] = list(set(analysis["placeholders_found"]))

            # 根据占位符建议字段
            analysis["suggested_fields"] = self._suggest_fields(analysis["placeholders_found"])

            # 根据文本内容建议字段
            text_based_fields = self._suggest_fields_from_text(analysis["extracted_text"])
            analysis["suggested_fields"].extend([f for f in text_based_fields if f not in analysis["suggested_fields"]])

            print(f"\n分析完成！")
            print(f"  找到占位符: {len(analysis['placeholders_found'])} 个")
            print(f"  建议字段: {len(analysis['suggested_fields'])} 个")
            print(f"  总占位符出现次数: {placeholder_count} 次")

            return analysis

        except Exception as e:
            print(f"分析模板失败: {str(e)}")
            import traceback
            traceback.print_exc()
            return {}

    def _find_placeholders(self, text: str) -> List[str]:
        """在文本中查找占位符"""
        placeholders = []

        # 检查各种占位符格式
        for i in range(0, len(self.placeholder_patterns), 2):
            start_pattern = self.placeholder_patterns[i]
            end_pattern = self.placeholder_patterns[i + 1]

            start_idx = text.find(start_pattern)
            while start_idx != -1:
                end_idx = text.find(end_pattern, start_idx + len(start_pattern))
                if end_idx != -1:
                    placeholder = text[start_idx + len(start_pattern):end_idx].strip()
                    if placeholder:
                        placeholders.append(placeholder)
                    start_idx = text.find(start_pattern, end_idx + len(end_pattern))
                else:
                    break

        return placeholders

    def _suggest_fields(self, placeholders: List[str]) -> List[Dict]:
        """根据占位符建议字段"""
        field_mapping = {
            # 名称相关
            "name": ["名称", "姓名", "name", "姓名", "经营者姓名", "个体工商户名称"],
            "business_name": ["个体工商户名称", "企业名称", "店铺名称", "字号", "商号"],
            "operator_name": ["经营者姓名", "负责人", "法人", "业主", "经营者"],

            # 证件相关
            "id_card": ["身份证", "证件号码", "身份证号", "身份证号码", "证件号"],
            "credit_code": ["统一社会信用代码", "信用代码", "社会信用代码", "组织机构代码"],

            # 联系信息
            "phone": ["电话", "联系电话", "手机", "手机号", "联系方式"],
            "email": ["邮箱", "电子邮箱", "email", "电子邮件"],
            "address": ["地址", "经营场所", "住所", "营业地址", "场所"],

            # 经营信息
            "business_scope": ["经营范围", "经营项目", "业务范围", "主营项目"],
            "registered_capital": ["资金", "注册资本", "出资额", "资金数额", "投资额"],
            "business_type": ["行业", "行业类型", "经营类型", "企业类型"],

            # 时间相关
            "application_date": ["日期", "申请日期", "提交日期", "办理日期"],
            "start_date": ["开业日期", "成立日期", "起始日期"],
            "end_date": ["截止日期", "结束日期", "有效期至"],
        }

        suggested_fields = []
        used_fields = set()

        for placeholder in placeholders:
            placeholder_lower = placeholder.lower()

            # 直接匹配
            for field_name, keywords in field_mapping.items():
                if field_name in used_fields:
                    continue

                for keyword in keywords:
                    if keyword in placeholder or keyword in placeholder_lower:
                        suggested_fields.append({
                            "field_name": field_name,
                            "display_name": keywords[0],
                            "placeholder": placeholder,
                            "description": f"对应占位符: {placeholder}",
                            "required": True
                        })
                        used_fields.add(field_name)
                        break

        # 添加标准字段（如果还没添加）
        standard_fields = [
            {
                "field_name": "business_name",
                "display_name": "个体工商户名称",
                "placeholder": "{{个体工商户名称}}",
                "description": "个体工商户的全称",
                "required": True
            },
            {
                "field_name": "operator_name",
                "display_name": "经营者姓名",
                "placeholder": "{{经营者姓名}}",
                "description": "经营者的姓名",
                "required": True
            },
            {
                "field_name": "id_card",
                "display_name": "身份证号码",
                "placeholder": "{{身份证号码}}",
                "description": "经营者的身份证号码",
                "required": True
            },
            {
                "field_name": "business_address",
                "display_name": "经营场所",
                "placeholder": "{{经营场所}}",
                "description": "个体工商户的经营地址",
                "required": True
            },
            {
                "field_name": "phone",
                "display_name": "联系电话",
                "placeholder": "{{联系电话}}",
                "description": "经营者的联系电话",
                "required": True
            },
            {
                "field_name": "business_scope",
                "display_name": "经营范围",
                "placeholder": "{{经营范围}}",
                "description": "个体工商户的经营范围",
                "required": True
            },
            {
                "field_name": "registered_capital",
                "display_name": "资金数额",
                "placeholder": "{{资金数额}}",
                "description": "个体工商户的资金数额（单位：元）",
                "required": True
            },
            {
                "field_name": "application_date",
                "display_name": "申请日期",
                "placeholder": "{{申请日期}}",
                "description": "申请日期（自动填充）",
                "required": False
            }
        ]

        for field in standard_fields:
            if field["field_name"] not in used_fields:
                suggested_fields.append(field)
                used_fields.add(field["field_name"])

        return suggested_fields

    def _suggest_fields_from_text(self, extracted_text: List[Dict]) -> List[Dict]:
        """从提取的文本中建议字段"""
        suggested_fields = []

        # 常见的关键词模式
        keyword_patterns = {
            "个体工商户名称": ["名称：", "名称为", "名称:", "个体工商户"],
            "经营者姓名": ["经营者：", "经营者姓名：", "负责人：", "法人："],
            "身份证号码": ["身份证：", "身份证号：", "证件号码：", "身份证号码："],
            "经营场所": ["场所：", "地址：", "经营场所：", "住所："],
            "经营范围": ["范围：", "经营范围：", "经营项目：", "业务范围："],
            "资金数额": ["资金：", "注册资本：", "出资额：", "资金数额："],
            "联系电话": ["电话：", "联系电话：", "手机：", "联系方式："]
        }

        for item in extracted_text:
            text = item["text"]
            for field_name, patterns in keyword_patterns.items():
                for pattern in patterns:
                    if pattern in text:
                        # 提取字段值（如果存在）
                        value_start = text.find(pattern) + len(pattern)
                        value_end = text.find("。", value_start) if "。" in text[value_start:] else len(text)
                        example_value = text[value_start:value_end].strip()

                        suggested_fields.append({
                            "field_name": field_name.lower().replace("：", "").replace(":", ""),
                            "display_name": field_name,
                            "placeholder": f"{{{{{field_name}}}}}",
                            "description": f"从模板中识别，示例值: {example_value[:20]}..." if example_value else f"从模板中识别",
                            "required": True
                        })
                        break

        return suggested_fields

    def generate_simple_filler(self, analysis: Dict, output_file: str = "simple_template_filler.py"):
        """生成简单的模板填充代码"""
        if not analysis:
            print("没有分析结果，无法生成代码")
            return

        template_name = analysis.get("template_name", "unknown_template.docx")
        suggested_fields = analysis.get("suggested_fields", [])

        code = f'''#!/usr/bin/env python3
"""
Word模板自动填充工具 - 简化版
专为模板: {template_name}
"""

import sys
from pathlib import Path
import json
import datetime

try:
    from docx import Document
except ImportError:
    print("需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)

class SimpleTemplateFiller:
    """{template_name} 模板填充器 - 简化版"""

    def __init__(self, template_path: str = "{template_name}"):
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            print(f"模板文件不存在: {{template_path}}")
            sys.exit(1)

    def fill_template(self, data: dict):
        """
        填充模板

        Args:
            data: 包含字段数据的字典
        """
        try:
            # 读取模板
            doc = Document(str(self.template_path))

            # 准备替换映射
            replacements = {{
'''

        # 添加字段映射
        for field in suggested_fields:
            field_name = field["field_name"]
            placeholder = field.get("placeholder", f"{{{{{field['display_name']}}}}}")
            code += f'                "{placeholder}": data.get("{field_name}", ""),\n'

        code += '''            }}

            # 添加自动生成的字段
            if "application_date" not in data:
                replacements["{{申请日期}}"] = datetime.datetime.now().strftime("%Y年%m月%d日")
            if "fill_date" not in data:
                replacements["{{填写日期}}"] = datetime.datetime.now().strftime("%Y年%m月%d日")

            # 在段落中替换
            for paragraph in doc.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, new_text)

            # 在表格中替换
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for old_text, new_text in replacements.items():
                            if old_text in cell.text:
                                cell.text = cell.text.replace(old_text, new_text)

            # 生成输出文件名
            business_name = data.get("business_name", "未知")
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = Path("filled_templates")
            output_dir.mkdir(exist_ok=True)
            output_path = output_dir / f"{business_name}_{timestamp}.docx"

            # 保存文件
            doc.save(str(output_path))

            print(f"模板填充成功！")
            print(f"文件保存位置: {{output_path}}")

            # 保存数据
            data_file = output_path.with_suffix('.json')
            with open(data_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            print(f"数据已保存: {{data_file}}")

            return output_path

        except Exception as e:
            print(f"填充模板失败: {{str(e)}}")
            return None

def collect_data():
    """收集数据"""
    print("请填写以下信息（按Enter跳过可选字段）：")
    data = {}

'''

        # 添加交互式数据收集代码
        for field in suggested_fields:
            field_name = field["field_name"]
            display_name = field["display_name"]
            required = field.get("required", True)

            if required:
                code += f'    data["{field_name}"] = input("{display_name}: ").strip()\n'
                code += f'    while not data["{field_name}"]:\n'
                code += f'        print("{display_name}是必填项")\n'
                code += f'        data["{field_name}"] = input("{display_name}: ").strip()\n'
            else:
                code += f'    value = input("{display_name}（可选）: ").strip()\n'
                code += f'    if value:\n'
                code += f'        data["{field_name}"] = value\n'
            code += '\n'

        code += '''    return data

def main():
    """主函数"""
    print("=" * 60)
    print(f"  {template_name} 模板填充工具")
    print("=" * 60)

    filler = SimpleTemplateFiller()

    print("\\n模板字段说明：")
'''

        # 添加字段说明
        for field in suggested_fields:
            required = "必填" if field.get("required", True) else "可选"
            code += f'    print(f"  - {{field[\"display_name\"]}} ({{required}}) - {{field[\"description\"]}}")'

        code += '''

    print("\\n1. 交互式填写")
    print("2. 从JSON文件加载")
    print("3. 退出")

    choice = input("\\n请选择操作 (1-3): ").strip()

    if choice == "1":
        data = collect_data()
    elif choice == "2":
        json_file = input("请输入JSON文件路径: ").strip()
        try:
            with open(json_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            print(f"从JSON文件加载 {{len(data)}} 个字段")
        except Exception as e:
            print(f"加载JSON文件失败: {{str(e)}}")
            return
    elif choice == "3":
        print("再见！")
        return
    else:
        print("无效选择")
        return

    # 填充模板
    filler.fill_template(data)

if __name__ == "__main__":
    main()
'''

        # 保存生成的代码
        output_path = Path(output_file)
        output_path.write_text(code, encoding='utf-8')

        print(f"\n已生成模板填充代码: {output_file}")
        print(f"使用方法:")
        print(f"  1. 确保模板文件 '{template_name}' 在当前目录")
        print(f"  2. 运行: python {output_file}")
        print(f"  3. 按照提示填写信息")

def main():
    """主函数"""
    print("=" * 60)
    print("  Word模板分析工具 - 简化版")
    print("  专为个体工商户登记申请书设计")
    print("=" * 60)

    analyzer = SimpleTemplateAnalyzer()

    # 查找模板文件
    template_files = list(Path(".").glob("*.docx"))
    template_files.extend(list(Path(".").glob("*.doc")))

    if not template_files:
        print("未找到Word模板文件")
        print("提示：请将Word模板文件放在当前目录")
        return

    print("\n找到以下模板文件：")
    for i, template_file in enumerate(template_files, 1):
        size_kb = template_file.stat().st_size / 1024
        print(f"{i}. {template_file.name} ({size_kb:.1f} KB)")

    try:
        choice = input("\n请选择要分析的模板文件编号: ").strip()
        if not choice:
            print("未选择，退出")
            return

        choice = int(choice)
        if 1 <= choice <= len(template_files):
            template_file = template_files[choice - 1]

            # 分析模板
            analysis = analyzer.analyze_template(template_file)

            if analysis:
                # 保存分析结果
                analysis_file = f"{template_file.stem}_analysis.json"
                with open(analysis_file, 'w', encoding='utf-8') as f:
                    json.dump(analysis, f, ensure_ascii=False, indent=2)
                print(f"\n分析结果已保存: {analysis_file}")

                # 显示分析结果摘要
                print("\n分析结果摘要：")
                print(f"  模板文件: {analysis['template_name']}")
                print(f"  文件大小: {analysis['template_size_kb']:.1f} KB")
                print(f"  段落数量: {analysis['paragraph_count']}")
                print(f"  表格数量: {analysis['table_count']}")
                print(f"  找到占位符: {len(analysis['placeholders_found'])} 个")

                if analysis['placeholders_found']:
                    print("  具体占位符：")
                    for placeholder in analysis['placeholders_found']:
                        print(f"    - {placeholder}")

                print(f"\n建议字段 ({len(analysis['suggested_fields'])} 个)：")
                for field in analysis['suggested_fields']:
                    required = "必填" if field.get("required", True) else "可选"
                    print(f"    - {field['display_name']} ({required}) - {field['description']}")

                # 询问是否生成填充代码
                generate_code = input("\n是否生成模板填充代码？ (y/n): ").strip().lower()
                if generate_code == 'y':
                    code_file = f"{template_file.stem}_filler.py"
                    analyzer.generate_simple_filler(analysis, code_file)

                    print(f"\n完成！")
                    print(f"  1. 分析结果: {analysis_file}")
                    print(f"  2. 填充代码: {code_file}")
                    print(f"\n下一步：")
                    print(f"  运行: python {code_file}")
                    print(f"  按照提示填写信息，自动生成填充后的Word文档")

        else:
            print("无效的选择")
    except ValueError:
        print("请输入有效的数字")
    except Exception as e:
        print(f"分析过程中出错: {str(e)}")

if __name__ == "__main__":
    main()