#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析修改后的模板，提取所有字段信息
"""

from docx import Document
import json

def get_color_info(run):
    """获取颜色信息"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        return f'RGB({r},{g},{b})'
    return 'default'

def analyze_template():
    """分析模板"""
    doc = Document(r'c:\Users\flyskyson\Desktop\（李奕凤）个体工商户开业登记申请书（模板-待修改）.docx')

    result = {
        'red_fields': [],      # 红色字段（需要填写的变量）
        'green_fields': [],    # 绿色字段（常量）
        'black_fields': [],    # 黑色字段（字段名）
        'all_fields': []       # 所有字段
    }

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text.strip()
                        if not text:
                            continue

                        color = get_color_info(run)

                        field_info = {
                            'table': table_idx,
                            'row': row_idx,
                            'cell': cell_idx,
                            'text': text,
                            'color': color
                        }

                        result['all_fields'].append(field_info)

                        # 分类
                        if 'RGB(255,0,0)' in color:
                            result['red_fields'].append(field_info)
                        elif 'RGB(0,128,0)' in color or 'RGB(0,255,0)' in color:
                            result['green_fields'].append(field_info)
                        elif color == 'default':
                            result['black_fields'].append(field_info)

    # 保存结果
    output_file = r'c:\Users\flyskyson\Desktop\template_analysis.json'
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    # 打印摘要
    print("=" * 70)
    print("模板分析结果")
    print("=" * 70)
    print(f"\n红色字段（变量）: {len(result['red_fields'])} 个")
    print(f"绿色字段（常量）: {len(result['green_fields'])} 个")
    print(f"黑色字段（字段名）: {len(result['black_fields'])} 个")

    print("\n【绿色字段（常量）】")
    for field in result['green_fields']:
        print(f"  - {field['text']}")

    print("\n【红色字段（变量）】")
    for field in result['red_fields']:
        print(f"  - {field['text']}")

    print(f"\n详细分析已保存到: {output_file}")

    return result

if __name__ == "__main__":
    analyze_template()
