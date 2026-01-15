#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模板转换工具 - 将颜色标记模板转换为 Jinja2 模板
"""

from docx import Document
import json

def convert_template():
    """转换模板"""

    # 读取原模板
    src = Document(r'c:\Users\flyskyson\Desktop\（李奕凤）个体工商户开业登记申请书（模板-待修改）.docx')

    # 创建新模板
    dst = Document()

    # 复制标题
    for para in src.paragraphs:
        if para.text.strip():
            dst.add_paragraph(para.text, para.style)

    # 复制表格
    for table in src.tables:
        # 创建新表格
        dst_table = dst.add_table(len(table.rows), len(table.columns))

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                # 复制文本
                dst_table.rows[i].cells[j].text = cell.text

    # 保存新模板
    output_file = r'c:\Users\flyskyson\Desktop\（李奕凤）个体工商户开业登记申请书（Jinja2模板-待编辑）.docx'
    dst.save(output_file)

    print("=" * 70)
    print("模板转换完成！")
    print("=" * 70)
    print(f"\n输出文件: {output_file}")
    print("\n下一步操作:")
    print("1. 在WPS中打开这个新模板")
    print("2. 找到原来的红色字体内容")
    print("3. 将它们改为 {{变量名}} 格式")
    print("4. 保存后使用 jinja2_filler.py 生成文档")
    print("\n变量名对照表:")
    print("-" * 70)

    # 显示变量映射
    mappings = {
        "个体工商户名称": "{{business_name}}",
        "经营者姓名": "{{operator_name}}",
        "经营者": "{{operator_name}}",
        "联系电话": "{{phone}}",
        "电话": "{{phone}}",
        "电子邮箱": "{{email}}",
        "邮箱": "{{email}}",
        "经营场所": "{{business_address}}",
        "住所": "{{business_address}}",
        "地址": "{{business_address}}",
        "邮政编码": "{{postal_code}}",
        "邮编": "{{postal_code}}",
        "从业人数": "{{employee_count}}",
        "人数": "{{employee_count}}",
        "注册资金": "{{registered_capital}}",
        "资金": "{{registered_capital}}",
        "身份证号码": "{{id_card}}",
        "身份证": "{{id_card}}",
        "性别": "{{gender}}",
        "民族": "{{nation}}",
        "政治面貌": "{{political_status}}",
        "政治": "{{political_status}}",
        "文化程度": "{{education}}",
        "学历": "{{education}}",
        "经营范围": "{{business_scope}}",
        "经营期限": "{{operation_period}}",
        "期限": "{{operation_period}}"
    }

    for cn, var in mappings.items():
        print(f"  {cn:12} → {var}")

    print("\n" + "=" * 70)

if __name__ == "__main__":
    convert_template()
