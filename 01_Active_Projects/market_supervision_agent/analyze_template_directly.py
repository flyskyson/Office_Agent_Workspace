#!/usr/bin/env python3
"""
直接分析Word模板
"""

import sys
from pathlib import Path
import json

try:
    from docx import Document
except ImportError:
    print("需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)

def analyze_template_directly():
    """直接分析模板"""
    template_name = "个体工商户开业登记申请书（模板）.docx"
    template_path = Path(template_name)

    if not template_path.exists():
        print(f"模板文件不存在: {template_name}")
        print("请确保文件在当前目录")
        return

    print(f"正在分析模板: {template_name}")
    print("文件大小: {:.1f} KB".format(template_path.stat().st_size / 1024))
    print("=" * 60)

    try:
        doc = Document(str(template_path))

        # 收集所有文本
        all_text = []

        # 收集段落文本
        print("\n段落分析:")
        for i, paragraph in enumerate(doc.paragraphs[:15]):  # 只分析前15个段落
            text = paragraph.text.strip()
            if text:
                all_text.append(text)
                print(f"  段落 {i+1}: {text[:50]}..." if len(text) > 50 else f"  段落 {i+1}: {text}")

        # 收集表格文本
        print(f"\n表格分析 (共 {len(doc.tables)} 个表格):")
        for i, table in enumerate(doc.tables):
            print(f"  表格 {i+1}: {len(table.rows)} 行 x {len(table.columns)} 列")
            for row_idx, row in enumerate(table.rows[:3]):  # 只显示前3行
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    print(f"    第{row_idx+1}行: {' | '.join(row_text)}")

        # 查找可能的占位符
        print("\n查找可能的字段位置:")
        field_keywords = {
            "个体工商户名称": ["名称", "个体工商户", "字号"],
            "经营者姓名": ["经营者", "姓名", "负责人"],
            "身份证号码": ["身份证", "证件", "号码"],
            "经营场所": ["场所", "地址", "经营地址"],
            "经营范围": ["范围", "经营项目", "业务"],
            "资金数额": ["资金", "注册资本", "出资"],
            "联系电话": ["电话", "联系", "手机"],
            "申请日期": ["日期", "申请", "提交"]
        }

        found_fields = {}
        for field_name, keywords in field_keywords.items():
            for text in all_text:
                for keyword in keywords:
                    if keyword in text:
                        if field_name not in found_fields:
                            found_fields[field_name] = []
                        found_fields[field_name].append(text[:100])
                        break

        print("\n找到的字段位置:")
        for field_name, examples in found_fields.items():
            print(f"  {field_name}:")
            for example in examples[:2]:  # 只显示前2个例子
                print(f"    - {example}")

        # 生成简单的填充程序
        generate_filler_code(found_fields, template_name)

        print("\n分析完成！")

    except Exception as e:
        print(f"分析失败: {str(e)}")
        import traceback
        traceback.print_exc()

def generate_filler_code(found_fields, template_name):
    """生成填充代码"""
    code = '''#!/usr/bin/env python3
"""
个体工商户开业登记申请书填充工具
"""

import sys
from pathlib import Path
import json
import datetime

try:
    from docx import Document
except ImportError:
    print("需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)

def fill_template():
    """填充模板"""
    template_file = "{}"
    if not Path(template_file).exists():
        print(f"模板文件不存在: {{template_file}}")
        return

    print("请填写以下信息:")
    data = {{}}

'''.format(template_name)

    # 添加字段输入
    field_mapping = {
        "个体工商户名称": "business_name",
        "经营者姓名": "operator_name",
        "身份证号码": "id_card",
        "经营场所": "business_address",
        "经营范围": "business_scope",
        "资金数额": "registered_capital",
        "联系电话": "phone",
        "申请日期": "application_date"
    }

    for display_name, field_name in field_mapping.items():
        if display_name in found_fields:
            code += f'    data["{field_name}"] = input("{display_name}: ").strip()\n'
        else:
            code += f'    data["{field_name}"] = input("{display_name} (可选): ").strip()\n'

    code += '''
    # 自动填充日期
    if not data.get("application_date"):
        data["application_date"] = datetime.datetime.now().strftime("%Y年%m月%d日")

    try:
        # 读取模板
        doc = Document(template_file)

        # 简单的文本替换
        replacements = {
'''

    for display_name, field_name in field_mapping.items():
        code += f'            "{{{display_name}}}": data.get("{field_name}", ""),\n'

    code += '''        }

        # 在段落中替换
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)

        # 在表格中替换
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old_text, new_text in replacements.items():
                        if old_text in cell.text:
                            cell.text = cell.text.replace(old_text, new_text)

        # 保存文件
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = f"filled_{{business_name}}_{{timestamp}}.docx"

        doc.save(output_file)

        print(f"\\n填充成功！")
        print(f"文件已保存: {{output_file}}")

        # 保存数据
        data_file = output_file.replace(".docx", ".json")
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"数据已保存: {{data_file}}")

    except Exception as e:
        print(f"填充失败: {{str(e)}}")

if __name__ == "__main__":
    print("个体工商户开业登记申请书填充工具")
    print("=" * 50)
    fill_template()
'''

    output_file = "simple_filler.py"
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(code)

    print(f"\n已生成填充程序: {output_file}")
    print("使用方法:")
    print(f"  1. 确保模板文件 '{template_name}' 在当前目录")
    print(f"  2. 运行: python {output_file}")
    print(f"  3. 按照提示填写信息")

if __name__ == "__main__":
    analyze_template_directly()