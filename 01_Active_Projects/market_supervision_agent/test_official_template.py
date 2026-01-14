#!/usr/bin/env python3
"""
测试填充官方模板
"""

import sys
import json
import datetime
from pathlib import Path
from docx import Document

def find_and_fill_in_runs(element, placeholder, new_text):
    """在runs中查找并替换，保持格式"""
    found = False

    # 遍历所有段落
    for paragraph in element.paragraphs:
        # 检查是否包含占位符
        if placeholder in paragraph.text:
            # 在runs中查找并替换
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, new_text)
                    found = True
                    break

    return found

def fill_official_template(data, template_file="个体工商户开业登记申请书（模板）.docx"):
    """填充官方模板"""
    try:
        print(f"正在读取模板: {template_file}")
        doc = Document(template_file)

        print("正在填充模板（保持原有格式）...")

        # 定义占位符映射
        placeholders = {
            "{{个体工商户名称}}": data.get("business_name", ""),
            "{{经营者姓名}}": data.get("operator_name", ""),
            "{{联系电话}}": data.get("phone", ""),
            "{{电子邮箱}}": data.get("email", ""),
            "{{经营场所}}": data.get("business_address", ""),
            "{{邮政编码}}": data.get("postal_code", ""),
            "{{从业人数}}": data.get("employee_count", ""),
            "{{注册资金}}": data.get("registered_capital", ""),
            "{{身份证号码}}": data.get("id_card", ""),
            "{{性别}}": data.get("gender", ""),
            "{{民族}}": data.get("nation", ""),
            "{{政治面貌}}": data.get("political_status", ""),
            "{{文化程度}}": data.get("education", ""),
            "{{经营范围}}": data.get("business_scope", ""),
            "{{经营期限}}": data.get("operation_period", ""),
            "{{申请日期}}": data.get("application_date", ""),
            "{{签字日期}}": data.get("sign_date", "")
        }

        replaced_count = 0

        # 在段落中替换
        for paragraph in doc.paragraphs:
            for placeholder, value in placeholders.items():
                if placeholder in paragraph.text:
                    find_and_fill_in_runs(paragraph, placeholder, value)
                    replaced_count += 1

        # 在表格中替换
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in placeholders.items():
                        if placeholder in cell.text:
                            find_and_fill_in_runs(cell, placeholder, value)
                            replaced_count += 1

        print(f"完成 {replaced_count} 处替换")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_官方模板_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] {data_file}")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数"""
    # 测试数据
    test_data = {
        "business_name": "王五快餐店",
        "operator_name": "王五",
        "phone": "13800138001",
        "email": "wangwu@example.com",
        "business_address": "广州市天河区天河路123号",
        "postal_code": "510000",
        "employee_count": "8",
        "registered_capital": "200000",
        "id_card": "440101198505056789",
        "gender": "男",
        "nation": "汉族",
        "political_status": "群众",
        "education": "大专",
        "business_scope": "快餐服务；餐饮配送",
        "operation_period": "长期",
        "application_date": datetime.datetime.now().strftime("%Y年%m月%d日"),
        "sign_date": datetime.datetime.now().strftime("%Y年%m月%d日")
    }

    print("=" * 60)
    print("测试填充官方模板")
    print("=" * 60)

    result = fill_official_template(test_data)

    if result:
        print("\n" + "=" * 60)
        print("测试成功！")
        print(f"生成的文件: {result}")
        print("=" * 60)
    else:
        print("\n测试失败！")

if __name__ == "__main__":
    main()
