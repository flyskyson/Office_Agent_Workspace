#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
李奕凤版个体工商户开业登记申请书填充工具 v2.0
- 支持修改后的模板
- 绿色字体（常量）保留不删除
- 红色字体（变量）删除并填充
- 支持全局配置文件
"""

import json
import datetime
import sys
import argparse
from pathlib import Path
from docx import Document

# ============ 配置加载 ============

def load_config(config_file="config.json"):
    """加载全局配置文件"""
    config_path = Path(__file__).parent / config_file
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except:
        # 如果配置文件不存在，返回默认配置
        return {
            'constants': {'postal_code_default': '537820'},
            'defaults': {'operation_period': '长期'},
            'field_mappings': {}
        }

# ============ 核心功能 ============

def is_red_font(run):
    """判断是否为红色字体 RGB(255,0,0)"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        return r == 255 and g == 0 and b == 0
    return False

def is_green_font(run):
    """判断是否为绿色字体 RGB(0,176,80)"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        return r == 0 and g == 176 and b == 80
    return False

def cell_has_font(cell, font_check_func):
    """检查单元格是否有特定颜色的字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if font_check_func(run):
                return True
    return False

def get_cell_text_clean(cell):
    """获取单元格的纯文本（用于字段名识别）"""
    text_parts = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if not is_red_font(run) and not is_green_font(run) and run.text.strip():
                text_parts.append(run.text.strip())
    return "".join(text_parts)

def fill_template(data, template_file="（李奕凤）个体工商户开业登记申请书（模板-待修改）.docx",
                  output_dir="output", auto_open=True, config=None):
    """
    填充李奕凤版模板 v2.0

    参数:
        data: 包含申请信息的字典
        template_file: 模板文件路径
        output_dir: 输出目录
        auto_open: 是否自动打开文档
        config: 全局配置

    返回:
        生成的文件路径，失败返回None
    """
    try:
        if config is None:
            config = load_config()

        print(f"正在读取模板: {template_file}")
        doc = Document(template_file)

        print("正在智能填充字段...")
        print("  - 红色字体（变量）将被替换")
        print("  - 绿色字体（常量）将保留")

        # 获取字段映射
        field_mappings = config.get('field_mappings', {})
        if not field_mappings:
            # 使用默认映射
            field_mappings = {
                "个体工商户名称": "business_name",
                "经营者姓名": "operator_name",
                "经营者": "operator_name",
                "联系电话": "phone",
                "电话": "phone",
                "电子邮箱": "email",
                "邮箱": "email",
                "经营场所": "business_address",
                "住所": "business_address",
                "地址": "business_address",
                "邮政编码": "postal_code",
                "邮编": "postal_code",
                "从业人数": "employee_count",
                "人数": "employee_count",
                "注册资金": "registered_capital",
                "资金": "registered_capital",
                "身份证号码": "id_card",
                "身份证": "id_card",
                "性别": "gender",
                "民族": "nation",
                "政治面貌": "political_status",
                "政治": "political_status",
                "文化程度": "education",
                "学历": "education",
                "经营范围": "business_scope",
                "经营期限": "operation_period",
                "期限": "operation_period"
            }

        deleted_count = 0
        filled_count = 0
        filled_fields = set()
        green_preserved = 0

        # 遍历所有表格
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # 只处理有红色字体的单元格
                    if not cell_has_font(cell, is_red_font):
                        continue

                    # 统计绿色字体（保留）
                    green_count = sum(1 for p in cell.paragraphs
                                     for r in p.runs if is_green_font(r))
                    if green_count > 0:
                        green_preserved += green_count

                    # 查找字段名
                    field_name_text = ""

                    # 优先左侧单元格
                    if cell_idx > 0:
                        left_cell = row.cells[cell_idx - 1]
                        field_name_text = get_cell_text_clean(left_cell)

                    # 删除红色字体（保留绿色和黑色）
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if is_red_font(run):
                                run.text = ""
                                deleted_count += 1

                    # 匹配并填充
                    matched_key = None
                    for field_name, data_key in field_mappings.items():
                        if data_key not in filled_fields:
                            if field_name in field_name_text:
                                matched_key = data_key
                                filled_fields.add(data_key)
                                break

                    if matched_key and matched_key in data:
                        value = str(data[matched_key])
                        if value and cell.paragraphs:
                            # 找到第一个非空段落或创建新段落
                            para = None
                            for p in cell.paragraphs:
                                if any(r.text.strip() for r in p.runs):
                                    para = p
                                    break
                            if not para and cell.paragraphs:
                                para = cell.paragraphs[0]

                            if para:
                                # 清空红色字体内容（保留绿色和黑色）
                                new_runs = []
                                for run in para.runs:
                                    if not is_red_font(run):
                                        new_runs.append(run)

                                para.text = ""
                                para.clear()

                                # 恢复非红色内容
                                for run in new_runs:
                                    para.add_run(run.text)

                                # 添加新数据
                                new_run = para.add_run(value)
                                new_run.font.name = '宋体'
                                new_run.font.size = 209712

                                filled_count += 1
                                print(f"  [OK] {matched_key}: {value}")

        # 保存文件
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        output_file = output_path / f"{business_name}_开业登记_{timestamp}.docx"

        doc.save(str(output_file))

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")
        print(f"[数据] {data_file}")
        print(f"[统计] 删除红色{deleted_count}处，填充{filled_count}处，保留绿色{green_preserved}处")

        # 自动打开
        if auto_open:
            import subprocess
            subprocess.Popen(['cmd', '/c', 'start', '', str(output_file)], shell=True)
            print("\n正在用WPS打开文档...")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

# ============ 主程序 ============

def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='李奕凤版个体工商户开业登记申请书填充工具 v2.0',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  # 使用测试数据
  python 李奕凤版申请书填充工具_v2.py --test

  # 从JSON文件加载数据
  python 李奕凤版申请书填充工具_v2.py --data test_data.json

  # 批量处理
  python 李奕凤版申请书填充工具_v2.py --batch data_list.json
        """
    )

    parser.add_argument('--data', type=str, help='JSON数据文件路径')
    parser.add_argument('--batch', type=str, help='批量处理的JSON文件路径')
    parser.add_argument('--template', type=str,
                        default='（李奕凤）个体工商户开业登记申请书（模板-待修改）.docx',
                        help='模板文件路径')
    parser.add_argument('--output', type=str, default='output', help='输出目录')
    parser.add_argument('--test', action='store_true', help='使用测试数据')
    parser.add_argument('--no-open', action='store_true', help='不自动打开文档')

    args = parser.parse_args()

    print("=" * 70)
    print("李奕凤版个体工商户开业登记申请书填充工具 v2.0")
    print("支持修改后的模板 - 绿色常量保留，红色变量填充")
    print("=" * 70)

    # 加载配置
    config = load_config()

    # 检查模板文件
    template_path = Path(args.template)
    if not template_path.is_absolute():
        template_path = Path(__file__).parent / args.template

    if not template_path.exists():
        print(f"\n[错误] 模板文件不存在: {template_path}")
        print("\n提示: 请将模板文件放在当前目录，或使用 --template 参数指定")
        return 1

    auto_open = not args.no_open

    # 根据参数执行不同操作
    if args.test:
        # 使用测试数据
        test_data = {
            'business_name': '测试便利店',
            'operator_name': '测试员',
            'phone': '13800138000',
            'email': 'test@example.com',
            'business_address': '广西玉林市兴业县蒲塘镇测试路123号',
            'postal_code': '537820',
            'employee_count': '2',
            'registered_capital': '30000',
            'id_card': '450101199001011234',
            'gender': '男',
            'nation': '汉族',
            'political_status': '群众',
            'education': '高中',
            'business_scope': '食品销售；日用百货',
            'operation_period': '长期'
        }

        print("\n[测试模式] 使用预设测试数据\n")
        result = fill_template(test_data, str(template_path), args.output, auto_open, config)

    elif args.batch:
        # 批量处理
        try:
            with open(args.batch, 'r', encoding='utf-8') as f:
                data_list = json.load(f)

            if isinstance(data_list, dict):
                data_list = [data_list]

            print(f"\n[批量处理] 共 {len(data_list)} 条数据\n")

            results = []
            for idx, data in enumerate(data_list, 1):
                print(f"\n处理第 {idx}/{len(data_list)} 条:")
                result = fill_template(data, str(template_path), args.output, False, config)
                if result:
                    results.append(result)

            # 全部处理完后，打开最后一个文件
            if auto_open and results:
                import subprocess
                subprocess.Popen(['cmd', '/c', 'start', '', str(results[-1])], shell=True)

            print(f"\n[完成] 共生成 {len(results)} 个文档")

        except Exception as e:
            print(f"\n[错误] 批量处理失败: {e}")
            return 1

    elif args.data:
        # 从文件加载
        try:
            with open(args.data, 'r', encoding='utf-8') as f:
                data = json.load(f)

            print(f"\n[加载数据] 从文件: {args.data}\n")
            result = fill_template(data, str(template_path), args.output, auto_open, config)

        except Exception as e:
            print(f"\n[错误] 加载数据失败: {e}")
            return 1

    else:
        # 显示帮助
        parser.print_help()
        print("\n快速开始:")
        print("  python 李奕凤版申请书填充工具_v2.py --test")
        return 0

    print("\n" + "=" * 70)
    print("处理完成！")
    print("=" * 70)
    return 0

if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        print("\n\n[中断] 用户取消操作")
        sys.exit(0)
    except Exception as e:
        print(f"\n[异常] {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
