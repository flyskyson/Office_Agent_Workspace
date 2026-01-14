#!/usr/bin/env python3
"""
智能填充李奕凤版模板 v6 - 简化版
- 删除所有红色字体内容
- 在同一单元格中添加实际数据
"""

import json
import datetime
from pathlib import Path
from docx import Document

def is_red_font(run):
    """判断是否为红色字体 RGB(255,0,0)"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        return r == 255 and g == 0 and b == 0
    return False

def fill_liyifeng_template_v6(data, template_file="（李奕凤）个体工商户开业登记申请书（模板）.docx"):
    """智能填充李奕凤版模板 v6"""
    try:
        print(f"正在读取模板: {template_file}")
        doc = Document(template_file)

        print("正在处理：删除红色字体，添加实际数据...")

        # 定义数据列表（按常见顺序排列）
        data_list = [
            ("个体工商户名称", data.get("business_name", "")),
            ("经营者", data.get("operator_name", "")),
            ("联系电话", data.get("phone", "")),
            ("电子邮箱", data.get("email", "")),
            ("经营场所", data.get("business_address", "")),
            ("邮政编码", data.get("postal_code", "")),
            ("从业人数", data.get("employee_count", "")),
            ("注册资金", data.get("registered_capital", "")),
            ("身份证", data.get("id_card", "")),
            ("性别", data.get("gender", "")),
            ("民族", data.get("nation", "")),
            ("政治面貌", data.get("political_status", "")),
            ("文化程度", data.get("education", "")),
            ("经营范围", data.get("business_scope", "")),
            ("经营期限", data.get("operation_period", ""))
        ]

        deleted_count = 0
        filled_count = 0
        filled_fields = []

        # 遍历所有表格
        for table_idx, table in enumerate(doc.tables):
            print(f"\n处理表格 {table_idx}...")

            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # 获取单元格的完整文本（包括红色和黑色）
                    cell_text = cell.text

                    # 检查单元格中是否有红色字体
                    has_red = False
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if is_red_font(run):
                                has_red = True
                                break
                        if has_red:
                            break

                    # 如果有红色字体
                    if has_red:
                        # 1. 先删除所有红色字体的内容
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if is_red_font(run):
                                    run.text = ""
                                    deleted_count += 1

                        # 2. 检查这个单元格应该填充什么数据
                        cell_text_after_delete = cell.text  # 删除红色后的文本

                        matched_data = None
                        for field_name, field_value in data_list:
                            if field_name in cell_text_after_delete and field_value:
                                matched_data = (field_name, field_value)
                                break

                        # 3. 如果匹配到数据，在单元格中添加
                        if matched_data:
                            field_name, field_value = matched_data

                            # 在单元格的最后一个段落添加数据
                            if cell.paragraphs:
                                last_para = cell.paragraphs[-1]

                                # 获取最后一个run的格式
                                if last_para.runs:
                                    last_run = last_para.runs[-1]

                                    # 添加新run，复制格式
                                    new_run = last_para.add_run(f"  {field_value}")
                                    if last_run.font.name:
                                        new_run.font.name = last_run.font.name
                                    if last_run.font.size:
                                        new_run.font.size = last_run.font.size

                                    filled_count += 1
                                    filled_fields.append(f"  [{field_name}] {field_value}")
                                    print(f"  [OK] {field_name}: {field_value}")

        print(f"\n完成统计:")
        print(f"  删除红色字体: {deleted_count} 处")
        print(f"  填充数据: {filled_count} 处")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_申请书_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")
        print(f"[说明] 已删除红色字体，添加实际数据")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] {data_file}")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数"""
    # 测试数据
    test_data = {
        'business_name': '杨十四服装店',
        'operator_name': '杨十四',
        'phone': '13999999999',
        'email': 'yangshisi@test.com',
        'business_address': '广州市天河区天河路388号',
        'postal_code': '510000',
        'employee_count': '10',
        'registered_capital': '250000',
        'id_card': '440101198801088888',
        'gender': '女',
        'nation': '汉族',
        'political_status': '党员',
        'education': '大专',
        'business_scope': '服装零售；鞋帽销售',
        'operation_period': '长期'
    }

    print("=" * 70)
    print("李奕凤版模板填充工具 v6 - 简化版")
    print("删除红色字体示例数据，在同一单元格添加实际数据")
    print("=" * 70)

    print("\n测试数据：")
    for key, value in test_data.items():
        print(f"  {key}: {value}")

    result = fill_liyifeng_template_v6(test_data)

    if result:
        print("\n" + "=" * 70)
        print("测试成功！")
        print(f"生成的文件: {result}")
        print("=" * 70)

        # 用WPS打开
        print("\n正在用WPS打开文档...")
        import subprocess
        subprocess.Popen(['cmd', '/c', 'start', '', str(result)], shell=True)

if __name__ == "__main__":
    main()
