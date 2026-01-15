#!/usr/bin/env python3
"""
改进版：智能填充官方模板
1. 去除"模板"等提示文字
2. 处理固定的测试数据（如518000、537820等）
"""

import sys
import json
import datetime
from pathlib import Path
from docx import Document

def find_cell_by_keyword(table, keywords):
    """根据关键字在表格中查找单元格"""
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            for keyword in keywords:
                if keyword in cell_text:
                    # 返回相邻的单元格（通常是下一列）
                    if cell_idx + 1 < len(row.cells):
                        return row.cells[cell_idx + 1]
    return None

def find_cell_by_value(table, value):
    """根据单元格内容查找单元格"""
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            if value in cell.text:
                return cell
    return None

def fill_cell_smart(cell, new_text):
    """智能填充单元格，保持原有格式"""
    if not cell or not cell.paragraphs:
        return

    # 在第一个段落的第一个run中填充，保持格式
    paragraph = cell.paragraphs[0]

    # 清除现有文本但保留runs
    for run in paragraph.runs:
        run.text = ""

    # 在第一个run中设置新文本
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        # 如果没有runs，添加新run
        run = paragraph.add_run(new_text)

def clean_template_text(doc):
    """清理模板中的提示文字"""
    template_texts = ["模板", "示例", "填写", "仅供参考"]
    cleaned_count = 0

    for paragraph in doc.paragraphs:
        for text in template_texts:
            if text in paragraph.text:
                paragraph.text = paragraph.text.replace(text, "")
                cleaned_count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for text in template_texts:
                    if text in cell.text:
                        for paragraph in cell.paragraphs:
                            paragraph.text = paragraph.text.replace(text, "")
                        cleaned_count += 1

    return cleaned_count

def fill_official_template_improved(data, template_file="个体工商户开业登记申请书（模板）.docx"):
    """改进的官方模板填充"""
    try:
        print(f"正在读取模板: {template_file}")
        doc = Document(template_file)

        print("正在智能填充模板（保持原有格式和版式）...")

        # 定义字段映射：关键字 -> 数据字段
        field_mappings = [
            (["个体工商户名称", "名称"], "business_name"),
            (["经营者姓名", "经营者"], "operator_name"),
            (["联系电话", "电话"], "phone"),
            (["电子邮箱", "邮箱"], "email"),
            (["经营场所", "住所"], "business_address"),
            (["邮政编码", "邮编"], "postal_code"),
            (["从业人数", "人数"], "employee_count"),
            (["注册资金", "资金"], "registered_capital"),
            (["身份证号码", "身份证"], "id_card"),
            (["性别"], "gender"),
            (["民族"], "nation"),
            (["政治面貌"], "political_status"),
            (["文化程度", "学历"], "education"),
            (["经营范围"], "business_scope"),
            (["经营期限"], "operation_period")
        ]

        filled_count = 0

        # 在所有表格中查找并填充
        for table in doc.tables:
            for keywords, data_key in field_mappings:
                target_cell = find_cell_by_keyword(table, keywords)
                if target_cell and data_key in data:
                    # 清理数据中的非法字符
                    value = str(data.get(data_key, ""))
                    clean_value = ''.join(c for c in value if ord(c) < 0xD800 or ord(c) > 0xDFFF)

                    fill_cell_smart(target_cell, clean_value)
                    filled_count += 1
                    print(f"  已填充: {keywords[0]} -> {clean_value}")

        # 清理模板文字
        print("\n正在清理模板提示文字...")
        cleaned = clean_template_text(doc)
        print(f"  已清理 {cleaned} 处模板文字")

        # 替换固定的测试数据
        print("\n正在替换固定测试数据...")
        replacements = {
            "518000": data.get("postal_code", ""),
            "537820": "",  # 这个留空
            "模板": "",
            "示例": ""
        }

        replacement_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    for old, new in replacements.items():
                        if old in cell_text:
                            # 找到并替换
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if old in run.text:
                                        if new:  # 如果有新值
                                            run.text = run.text.replace(old, new)
                                            print(f"  已替换: {old} -> {new}")
                                        else:  # 如果新值为空，删除这部分
                                            run.text = run.text.replace(old, "")
                                            print(f"  已清除: {old}")
                                        replacement_count += 1

        print(f"\n完成统计:")
        print(f"  字段填充: {filled_count} 处")
        print(f"  模板清理: {cleaned} 处")
        print(f"  数据替换: {replacement_count} 处")
        print(f"  总计: {filled_count + cleaned + replacement_count} 处")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_开业登记申请书_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")
        print(f"[说明] 保持原有格式、版式和所有特殊符号")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] {data_file}")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def collect_data():
    """收集用户数据"""
    print("\n请填写个体工商户开业登记信息")
    print("=" * 60)
    print("说明: 带*的为必填项，其他可以按Enter跳过")

    data = {}

    # 基本信息
    print("\n[基本信息]")
    data["business_name"] = input("*个体工商户名称: ").strip()
    data["operator_name"] = input("*经营者姓名: ").strip()
    data["phone"] = input("*联系电话: ").strip()
    data["email"] = input("  电子邮箱（可选）: ").strip() or ""
    data["business_address"] = input("*经营场所: ").strip()
    data["postal_code"] = input("*邮政编码（可选）: ").strip() or ""
    data["employee_count"] = input("  从业人数（可选）: ").strip() or ""
    data["registered_capital"] = input("*注册资金（元）: ").strip()

    # 经营者信息
    print("\n[经营者信息]")
    data["id_card"] = input("*身份证号码: ").strip()
    data["gender"] = input("*性别（男/女）: ").strip()
    data["nation"] = input("  民族（可选）: ").strip() or ""
    data["political_status"] = input("  政治面貌（可选）: ").strip() or ""
    data["education"] = input("  文化程度（可选）: ").strip() or ""

    # 经营信息
    print("\n[经营信息]")
    data["business_scope"] = input("*经营范围: ").strip()
    data["operation_period"] = input("*经营期限（留空默认长期）: ").strip() or "长期"

    return data

def main():
    """主函数"""
    print("=" * 60)
    print("个体工商户开业登记申请书填充工具（官方格式-改进版）")
    print("智能识别字段，保持原有格式和版式")
    print("=" * 60)

    # 收集数据
    data = collect_data()

    # 填充模板
    result = fill_official_template_improved(data)

    if result:
        print("\n" + "=" * 60)
        print("完成！")
        print("请在WPS中打开生成的文件检查")
        print("=" * 60)

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n程序已取消")
    except Exception as e:
        print(f"\n程序错误: {str(e)}")
