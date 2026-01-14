#!/usr/bin/env python3
"""
测试用户修改后的模板
"""

import json
import datetime
from pathlib import Path
from docx import Document

def find_cell_by_keyword(table, keywords):
    """根据关键字在表格中查找单元格"""
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            for keyword in keywords:
                if keyword in cell_text:
                    # 返回相邻的单元格（通常是下一列）
                    if cell_idx + 1 < len(row.cells):
                        return row.cells[cell_idx + 1]
    return None

def fill_cell_smart(cell, new_text):
    """智能填充单元格，保持原有格式"""
    if not cell or not cell.paragraphs:
        return

    # 在第一个段落的第一个run中填充，保持格式
    paragraph = cell.paragraphs[0]

    # 清除现有文本但保留runs
    for run in paragraph.runs:
        run.text = ""

    # 在第一个run中设置新文本
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        # 如果没有runs，添加新run
        run = paragraph.add_run(new_text)

def fill_template_with_user_version(data, template_file="个体工商户开业登记申请书_用户修改版.docx"):
    """使用用户修改后的模板填充"""
    try:
        print(f"正在读取用户修改后的模板: {template_file}")
        doc = Document(template_file)

        print("正在智能填充模板（保持原有格式和版式）...")

        # 定义字段映射：关键字 -> 数据字段
        field_mappings = [
            (["个体工商户名称", "名称"], "business_name"),
            (["经营者姓名", "经营者"], "operator_name"),
            (["联系电话", "电话"], "phone"),
            (["电子邮箱", "邮箱"], "email"),
            (["经营场所", "住所"], "business_address"),
            (["邮政编码", "邮编"], "postal_code"),
            (["从业人数", "人数"], "employee_count"),
            (["注册资金", "资金"], "registered_capital"),
            (["身份证号码", "身份证"], "id_card"),
            (["性别"], "gender"),
            (["民族"], "nation"),
            (["政治面貌"], "political_status"),
            (["文化程度", "学历"], "education"),
            (["经营范围"], "business_scope"),
            (["经营期限"], "operation_period")
        ]

        filled_count = 0

        # 在所有表格中查找并填充
        for table in doc.tables:
            for keywords, data_key in field_mappings:
                target_cell = find_cell_by_keyword(table, keywords)
                if target_cell and data_key in data:
                    # 清理数据中的非法字符
                    value = str(data.get(data_key, ""))
                    clean_value = ''.join(c for c in value if ord(c) < 0xD800 or ord(c) > 0xDFFF)

                    fill_cell_smart(target_cell, clean_value)
                    filled_count += 1
                    print(f"  已填充: {keywords[0]} -> {clean_value}")

        print(f"\n完成 {filled_count} 处填充")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_申请书_用户模板_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")
        print(f"[说明] 使用用户修改后的模板，保持所有格式")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] {data_file}")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数"""
    # 新的测试数据
    test_data = {
        'business_name': '周八文具店',
        'operator_name': '周八',
        'phone': '13555555555',
        'email': 'zhouba@test.com',
        'business_address': '成都市锦江区春熙路666号',
        'postal_code': '610000',
        'employee_count': '6',
        'registered_capital': '150000',
        'id_card': '510101198707076666',
        'gender': '男',
        'nation': '汉族',
        'political_status': '群众',
        'education': '大专',
        'business_scope': '文具零售；办公用品销售',
        'operation_period': '长期'
    }

    print("=" * 60)
    print("测试用户修改后的模板")
    print("=" * 60)

    print("\n测试数据：")
    for key, value in test_data.items():
        print(f"  {key}: {value}")

    result = fill_template_with_user_version(test_data)

    if result:
        print("\n" + "=" * 60)
        print("测试成功！")
        print(f"生成的文件: {result}")
        print("=" * 60)

        # 用WPS打开
        print("\n正在用WPS打开文档...")
        import subprocess
        subprocess.Popen(['cmd', '/c', 'start', '', str(result)], shell=True)

if __name__ == "__main__":
    main()
