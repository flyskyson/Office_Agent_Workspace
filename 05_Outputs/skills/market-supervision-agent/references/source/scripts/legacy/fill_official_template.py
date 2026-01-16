#!/usr/bin/env python3
"""
填充官方模板 - 保持原有格式和版式
"""

import sys
import json
import datetime
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement

def fill_cell_preserving_format(cell, new_text):
    """填充单元格但保持原有格式"""
    # 清理单元格内容但保留段落格式
    for paragraph in cell.paragraphs:
        # 清除所有runs
        for run in paragraph.runs:
            run.text = ""

    # 在第一个run中设置新文本（保持格式）
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = new_text
    else:
        # 如果没有run，创建一个新的
        if cell.paragraphs:
            cell.paragraphs[0].text = new_text
        else:
            cell.text = new_text

def find_and_fill_in_runs(element, placeholder, new_text):
    """在runs中查找并替换，保持格式"""
    found = False

    # 遍历所有段落
    for paragraph in element.paragraphs:
        # 检查是否包含占位符
        if placeholder in paragraph.text:
            # 在runs中查找并替换
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, new_text)
                    found = True
                    break

    return found

def fill_official_template(data, template_file="个体工商户开业登记申请书（模板）.docx"):
    """填充官方模板"""
    try:
        print(f"正在读取模板: {template_file}")
        doc = Document(template_file)

        print("正在填充模板（保持原有格式）...")

        # 定义占位符映射
        placeholders = {
            "{{个体工商户名称}}": data.get("business_name", ""),
            "{{经营者姓名}}": data.get("operator_name", ""),
            "{{联系电话}}": data.get("phone", ""),
            "{{电子邮箱}}": data.get("email", ""),
            "{{经营场所}}": data.get("business_address", ""),
            "{{邮政编码}}": data.get("postal_code", ""),
            "{{从业人数}}": data.get("employee_count", ""),
            "{{注册资金}}": data.get("registered_capital", ""),
            "{{身份证号码}}": data.get("id_card", ""),
            "{{性别}}": data.get("gender", ""),
            "{{民族}}": data.get("nation", ""),
            "{{政治面貌}}": data.get("political_status", ""),
            "{{文化程度}}": data.get("education", ""),
            "{{经营范围}}": data.get("business_scope", ""),
            "{{经营期限}}": data.get("operation_period", ""),
            "{{申请日期}}": data.get("application_date", ""),
            "{{签字日期}}": data.get("sign_date", "")
        }

        replaced_count = 0

        # 在段落中替换
        for paragraph in doc.paragraphs:
            for placeholder, value in placeholders.items():
                if placeholder in paragraph.text:
                    find_and_fill_in_runs(paragraph, placeholder, value)
                    replaced_count += 1

        # 在表格中替换
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in placeholders.items():
                        if placeholder in cell.text:
                            find_and_fill_in_runs(cell, placeholder, value)
                            replaced_count += 1

        print(f"完成 {replaced_count} 处替换")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_开业登记申请书_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] {data_file}")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def collect_data():
    """收集用户数据"""
    print("\n请填写个体工商户开业登记信息")
    print("=" * 60)
    print("说明: 带*的为必填项，其他可以按Enter跳过")

    data = {}

    # 基本信息
    print("\n[基本信息]")
    data["business_name"] = input("*个体工商户名称: ").strip()
    data["operator_name"] = input("*经营者姓名: ").strip()
    data["phone"] = input("*联系电话: ").strip()
    data["email"] = input("  电子邮箱（可选）: ").strip() or "未填写"
    data["business_address"] = input("*经营场所: ").strip()
    data["postal_code"] = input("  邮政编码（可选）: ").strip() or "未填写"
    data["employee_count"] = input("  从业人数（可选）: ").strip() or "未填写"
    data["registered_capital"] = input("*注册资金（元）: ").strip()

    # 经营者信息
    print("\n[经营者信息]")
    data["id_card"] = input("*身份证号码: ").strip()
    data["gender"] = input("*性别（男/女）: ").strip()
    data["nation"] = input("  民族（可选）: ").strip() or "未填写"
    data["political_status"] = input("  政治面貌（可选）: ").strip() or "未填写"
    data["education"] = input("  文化程度（可选）: ").strip() or "未填写"

    # 经营信息
    print("\n[经营信息]")
    data["business_scope"] = input("*经营范围: ").strip()
    data["operation_period"] = input("*经营期限: ").strip() or "长期"

    # 自动添加日期
    data["application_date"] = datetime.datetime.now().strftime("%Y年%m月%d日")
    data["sign_date"] = data["application_date"]

    return data

def main():
    """主函数"""
    print("=" * 60)
    print("个体工商户开业登记申请书填充工具")
    print("使用官方模板，保持原有格式和版式")
    print("=" * 60)

    # 收集数据
    data = collect_data()

    # 填充模板
    result = fill_official_template(data)

    if result:
        print("\n" + "=" * 60)
        print("完成！")
        print(f"生成的文件在 output/ 目录")
        print("=" * 60)

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n程序已取消")
    except Exception as e:
        print(f"\n程序错误: {str(e)}")
