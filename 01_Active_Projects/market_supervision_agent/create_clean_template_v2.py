#!/usr/bin/env python3
"""
创建干净的Word模板 - 修复编码问题
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_clean_template():
    """创建干净的模板"""
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)

    # 标题
    title = doc.add_heading('个体工商户开业登记申请书', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加空行
    doc.add_paragraph('')

    # 创建信息表格
    table = doc.add_table(rows=15, cols=4)
    table.style = 'Table Grid'

    # 设置表格宽度
    table.width = Inches(6)

    # 填充表格数据
    cells_data = [
        # (行, 列, 文本, 跨行, 跨列)
        (0, 0, '个体工商户名称', 1, 1),
        (0, 1, '{{个体工商户名称}}', 1, 3),

        (1, 0, '经营者姓名', 1, 1),
        (1, 1, '{{经营者姓名}}', 1, 1),
        (1, 2, '联系电话', 1, 1),
        (1, 3, '{{联系电话}}', 1, 1),

        (2, 0, '电子邮箱', 1, 1),
        (2, 1, '{{电子邮箱}}', 1, 1),
        (2, 2, '邮政编码', 1, 1),
        (2, 3, '{{邮政编码}}', 1, 1),

        (3, 0, '经营场所', 1, 1),
        (3, 1, '{{经营场所}}', 1, 3),

        (4, 0, '从业人数', 1, 1),
        (4, 1, '{{从业人数}}', 1, 1),
        (4, 2, '注册资金(元)', 1, 1),
        (4, 3, '{{注册资金}}', 1, 1),

        (5, 0, '身份证号码', 1, 1),
        (5, 1, '{{身份证号码}}', 1, 3),

        (6, 0, '性别', 1, 1),
        (6, 1, '{{性别}}', 1, 1),
        (6, 2, '民族', 1, 1),
        (6, 3, '{{民族}}', 1, 1),

        (7, 0, '政治面貌', 1, 1),
        (7, 1, '{{政治面貌}}', 1, 1),
        (7, 2, '文化程度', 1, 1),
        (7, 3, '{{文化程度}}', 1, 1),

        (8, 0, '经营范围', 1, 1),
        (8, 1, '{{经营范围}}', 1, 3),

        (9, 0, '经营期限', 1, 1),
        (9, 1, '{{经营期限}}', 1, 3),

        (10, 0, '申请日期', 1, 1),
        (10, 1, '{{申请日期}}', 1, 3),

        (11, 0, '', 1, 4),
        (12, 0, '', 1, 4),
        (13, 0, '', 1, 4),

        (14, 0, '经营者签字：', 1, 1),
        (14, 2, '{{签字日期}}', 1, 2),
    ]

    # 填充表格
    for row_idx, col_idx, text, rowspan, colspan in cells_data:
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = text

        # 设置字体
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(11)

    # 保存文件
    output_file = '个体工商户开业登记申请书_干净版.docx'
    doc.save(output_file)
    print(f'模板已创建: {output_file}')
    return output_file

if __name__ == '__main__':
    create_clean_template()
