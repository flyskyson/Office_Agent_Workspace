#!/usr/bin/env python3
"""
分析李奕凤版模板结构
"""

from docx import Document
from docx.shared import RGBColor

def analyze_template():
    """分析模板结构"""
    template_file = "（李奕凤）个体工商户开业登记申请书（模板）.docx"
    doc = Document(template_file)

    print("=" * 70)
    print(f"分析模板: {template_file}")
    print("=" * 70)

    # 分析表格
    print(f"\n表格数量: {len(doc.tables)}")

    for table_idx, table in enumerate(doc.tables):
        print(f"\n--- 表格 {table_idx} ---")
        print(f"行数: {len(table.rows)}, 列数: {len(table.columns)}")

        # 显示前10行内容
        for row_idx, row in enumerate(table.rows[:10]):
            row_info = []
            for cell_idx, cell in enumerate(row.cells):
                # 获取单元格文本
                cell_text = cell.text.strip()[:30]

                # 检查字体颜色
                colors = []
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.color and run.font.color.rgb:
                            r, g, b = run.font.color.rgb
                            if r > 200 and g < 100 and b < 100:
                                colors.append("红色")
                            elif r < 50 and g < 50 and b < 50:
                                colors.append("黑色")
                            else:
                                colors.append(f"RGB({r},{g},{b})")

                color_info = f" [{', '.join(colors)}]" if colors else ""
                row_info.append(f"{cell_text}{color_info}")

            print(f"  行{row_idx}: {' | '.join(row_info)}")

if __name__ == "__main__":
    analyze_template()
