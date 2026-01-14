#!/usr/bin/env python3
"""
分析官方模板结构
"""

from docx import Document
import json

def analyze_template():
    """分析模板结构"""
    template_file = "个体工商户开业登记申请书（模板）.docx"

    try:
        doc = Document(template_file)

        print("=" * 70)
        print(f"分析模板: {template_file}")
        print("=" * 70)

        # 分析段落
        print("\n[段落信息]")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"段落 {i}: {para.text[:50]}..." if len(para.text) > 50 else f"段落 {i}: {para.text}")

        # 分析表格
        print(f"\n[表格信息] 共 {len(doc.tables)} 个表格")
        for table_idx, table in enumerate(doc.tables):
            print(f"\n表格 {table_idx}:")
            print(f"  行数: {len(table.rows)}")
            print(f"  列数: {len(table.columns) if table.rows else 0}")

            # 显示前5行内容
            print(f"  内容预览:")
            for row_idx, row in enumerate(table.rows[:5]):
                row_text = " | ".join([cell.text.strip()[:20] for cell in row.cells])
                print(f"    行{row_idx}: {row_text}")

        # 保存分析结果
        analysis = {
            "paragraphs_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
            "template_file": template_file
        }

        with open("template_analysis.json", "w", encoding="utf-8") as f:
            json.dump(analysis, f, ensure_ascii=False, indent=2)

        print("\n分析完成！")

    except Exception as e:
        print(f"分析失败: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    analyze_template()
