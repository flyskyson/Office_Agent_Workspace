#!/usr/bin/env python3
"""
智能填充李奕凤版模板 v4 - 直接替换所有红色字体
"""

import json
import datetime
from pathlib import Path
from docx import Document

def is_red_font(run):
    """判断是否为红色字体 RGB(255,0,0)"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        return r == 255 and g == 0 and b == 0
    return False

def get_all_cell_text(cell):
    """获取单元格的所有文字（用于上下文匹配）"""
    all_text = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                all_text.append(run.text.strip())
    return " ".join(all_text)

def fill_liyifeng_template_v4(data, template_file="（李奕凤）个体工商户开业登记申请书（模板）.docx"):
    """智能填充李奕凤版模板 v4"""
    try:
        print(f"正在读取模板: {template_file}")
        doc = Document(template_file)

        print("正在智能填充红色字体字段...")

        # 定义字段到数据的映射
        field_data_map = {
            "business_name": data.get("business_name", ""),
            "operator_name": data.get("operator_name", ""),
            "phone": data.get("phone", ""),
            "email": data.get("email", ""),
            "business_address": data.get("business_address", ""),
            "postal_code": data.get("postal_code", ""),
            "employee_count": data.get("employee_count", ""),
            "registered_capital": data.get("registered_capital", ""),
            "id_card": data.get("id_card", ""),
            "gender": data.get("gender", ""),
            "nation": data.get("nation", ""),
            "political_status": data.get("political_status", ""),
            "education": data.get("education", ""),
            "business_scope": data.get("business_scope", ""),
            "operation_period": data.get("operation_period", "")
        }

        # 定义字段识别规则
        field_rules = [
            (["个体工商户名称", "名称"], "business_name"),
            (["经营者"], "operator_name"),
            (["联系电话", "电话"], "phone"),
            (["电子邮箱", "邮箱"], "email"),
            (["经营场所", "住所", "地址"], "business_address"),
            (["邮政编码", "邮编"], "postal_code"),
            (["从业人数", "人数"], "employee_count"),
            (["注册资金", "资金"], "registered_capital"),
            (["身份证"], "id_card"),
            (["性别"], "gender"),
            (["民族"], "nation"),
            (["政治面貌"], "political_status"),
            (["文化程度", "学历"], "education"),
            (["经营范围"], "business_scope"),
            (["经营期限", "期限"], "operation_period")
        ]

        filled_count = 0
        total_red_runs = 0

        # 遍历所有表格
        for table_idx, table in enumerate(doc.tables):
            print(f"\n处理表格 {table_idx}...")

            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # 获取单元格的所有文字作为上下文
                    cell_context = get_all_cell_text(cell)

                    # 检查单元格中是否有红色字体
                    has_red = False
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if is_red_font(run):
                                has_red = True
                                total_red_runs += 1
                                break
                        if has_red:
                            break

                    # 如果有红色字体，根据上下文决定填充什么数据
                    if has_red:
                        # 查找匹配的字段
                        matched_field = None
                        for keywords, field_key in field_rules:
                            for keyword in keywords:
                                if keyword in cell_context:
                                    matched_field = field_key
                                    break
                            if matched_field:
                                break

                        # 如果找到匹配，替换红色字体
                        if matched_field and field_data_map[matched_field]:
                            value = str(field_data_map[matched_field])
                            clean_value = ''.join(c for c in value if ord(c) < 0xD800 or ord(c) > 0xDFFF)

                            # 替换所有红色字体
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if is_red_font(run):
                                        old_text = run.text
                                        if old_text.strip():
                                            run.text = clean_value
                                            filled_count += 1
                                            print(f"  [OK] [{matched_field}] {old_text[:20]}... -> {clean_value}")

        print(f"\n完成统计:")
        print(f"  发现红色字体: {total_red_runs} 处")
        print(f"  成功填充: {filled_count} 处")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_李奕凤版_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")
        print(f"[说明] 李奕凤版模板，红色字体已替换为实际数据")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] {data_file}")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数"""
    # 测试数据
    test_data = {
        'business_name': '陈十二花店',
        'operator_name': '陈十二',
        'phone': '13111111111',
        'email': 'chenshier@test.com',
        'business_address': '重庆市渝中区解放碑步行街88号',
        'postal_code': '400000',
        'employee_count': '6',
        'registered_capital': '120000',
        'id_card': '500101198303032222',
        'gender': '女',
        'nation': '汉族',
        'political_status': '群众',
        'education': '高中',
        'business_scope': '花卉零售；花艺设计',
        'operation_period': '长期'
    }

    print("=" * 70)
    print("李奕凤版模板填充工具 v4")
    print("红色字体 RGB(255,0,0) = 示例数据（自动替换）")
    print("=" * 70)

    print("\n测试数据：")
    for key, value in test_data.items():
        print(f"  {key}: {value}")

    result = fill_liyifeng_template_v4(test_data)

    if result:
        print("\n" + "=" * 70)
        print("测试成功！")
        print(f"生成的文件: {result}")
        print("=" * 70)

        # 用WPS打开
        print("\n正在用WPS打开文档...")
        import subprocess
        subprocess.Popen(['cmd', '/c', 'start', '', str(result)], shell=True)

if __name__ == "__main__":
    main()
