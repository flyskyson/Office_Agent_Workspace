#!/usr/bin/env python3
"""
智能填充李奕凤版模板 v7 - 完美版
- 删除所有红色字体内容
- 根据单元格位置和上下文智能填充数据
"""

import json
import datetime
from pathlib import Path
from docx import Document

def is_red_font(run):
    """判断是否为红色字体 RGB(255,0,0)"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        return r == 255 and g == 0 and b == 0
    return False

def get_cell_text_without_color(cell):
    """获取单元格中非红色字体的文本（字段名）"""
    text_parts = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if not is_red_font(run) and run.text.strip():
                text_parts.append(run.text.strip())
    return "".join(text_parts)

def fill_liyifeng_template_v7(data, template_file="（李奕凤）个体工商户开业登记申请书（模板）.docx"):
    """智能填充李奕凤版模板 v7"""
    try:
        print(f"正在读取模板: {template_file}")
        doc = Document(template_file)

        print("正在处理：删除红色字体，智能填充数据...")

        # 定义字段识别规则（优先级从高到低）
        field_patterns = [
            # (字段名列表, 数据键, 位置提示)
            (["个体工商户名称"], "business_name", "first"),
            (["经营者姓名", "经营者"], "operator_name", "early"),
            (["联系电话", "联系", "电话"], "phone", "contact"),
            (["电子邮箱", "邮箱"], "email", "contact"),
            (["经营场所", "住所", "地址", "经营"], "business_address", "address"),
            (["邮政编码", "邮编"], "postal_code", "address"),
            (["从业人数", "从业", "人数"], "employee_count", "basic"),
            (["注册资金", "注册", "资金"], "registered_capital", "money"),
            (["身份证号码", "身份证"], "id_card", "id"),
            (["性别"], "gender", "personal"),
            (["民族"], "nation", "personal"),
            (["政治面貌", "政治"], "political_status", "personal"),
            (["文化程度", "学历"], "education", "personal"),
            (["经营范围"], "business_scope", "business"),
            (["经营期限", "期限"], "operation_period", "business")
        ]

        deleted_count = 0
        filled_count = 0

        # 记录已填充的字段，避免重复
        filled_fields = set()

        # 遍历所有表格
        for table_idx, table in enumerate(doc.tables):
            print(f"\n处理表格 {table_idx}...")

            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # 检查单元格中是否有红色字体
                    has_red = False
                    red_runs = []

                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        for run_idx, run in enumerate(paragraph.runs):
                            if is_red_font(run):
                                has_red = True
                                red_runs.append((para_idx, run_idx))

                    if has_red:
                        # 1. 删除所有红色字体内容
                        for para_idx, run_idx in red_runs:
                            try:
                                cell.paragraphs[para_idx].runs[run_idx].text = ""
                                deleted_count += 1
                            except:
                                pass

                        # 2. 获取删除后的文本（字段名）
                        field_name_text = get_cell_text_without_color(cell)

                        # 3. 查找匹配的字段
                        matched_key = None
                        for field_names, data_key, _ in field_patterns:
                            if data_key not in filled_fields:  # 未填充过的字段
                                for field_name in field_names:
                                    if field_name in field_name_text:
                                        matched_key = data_key
                                        filled_fields.add(data_key)
                                        break
                                if matched_key:
                                    break

                        # 4. 如果找到匹配，添加数据
                        if matched_key and matched_key in data:
                            value = str(data[matched_key])
                            if value:
                                # 在单元格的最后添加数据
                                if cell.paragraphs:
                                    last_para = cell.paragraphs[-1]

                                    # 获取格式
                                    font_name = None
                                    font_size = None
                                    if last_para.runs:
                                        last_run = last_para.runs[-1]
                                        if last_run.font.name:
                                            font_name = last_run.font.name
                                        if last_run.font.size:
                                            font_size = last_run.font.size

                                    # 添加新run
                                    new_run = last_para.add_run(f"  {value}")
                                    if font_name:
                                        new_run.font.name = font_name
                                    if font_size:
                                        new_run.font.size = font_size

                                    filled_count += 1
                                    print(f"  [OK] {matched_key}: {value}")

        print(f"\n完成统计:")
        print(f"  删除红色字体: {deleted_count} 处")
        print(f"  填充数据: {filled_count} 处")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_开业登记_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")
        print(f"[说明] 已删除红色字体示例数据，填充实际信息")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] {data_file}")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数"""
    # 测试数据
    test_data = {
        'business_name': '刘十五咖啡馆',
        'operator_name': '刘十五',
        'phone': '13888888888',
        'email': 'liushiwu@test.com',
        'business_address': '深圳市南山区科技园南区68号',
        'postal_code': '518000',
        'employee_count': '8',
        'registered_capital': '180000',
        'id_card': '440301198612069999',
        'gender': '男',
        'nation': '汉族',
        'political_status': '群众',
        'education': '本科',
        'business_scope': '咖啡服务；西点制作',
        'operation_period': '长期'
    }

    print("=" * 70)
    print("李奕凤版模板填充工具 v7 - 完美版")
    print("删除红色字体 + 智能填充数据")
    print("=" * 70)

    print("\n测试数据：")
    for key, value in test_data.items():
        print(f"  {key}: {value}")

    result = fill_liyifeng_template_v7(test_data)

    if result:
        print("\n" + "=" * 70)
        print("测试成功！")
        print(f"生成的文件: {result}")
        print("=" * 70)

        # 用WPS打开
        print("\n正在用WPS打开文档...")
        import subprocess
        subprocess.Popen(['cmd', '/c', 'start', '', str(result)], shell=True)

if __name__ == "__main__":
    main()
