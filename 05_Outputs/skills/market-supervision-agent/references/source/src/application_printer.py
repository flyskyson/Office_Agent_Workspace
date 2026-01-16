#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
申请书生成和打印工具 v5.0

功能：
- 生成Word申请书
- 转换为PDF
- 打印管理
- 打印记录

作者: Claude Code
日期: 2026-01-15
"""

import os
import sys
import json
from pathlib import Path
from typing import Dict, Optional, Any
from datetime import datetime
from dataclasses import dataclass, field
from loguru import logger

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docxtpl import DocxTemplate
    from jinja2 import Environment, BaseLoader
except ImportError:
    print("警告: python-docx 或 docxtpl 未安装，请运行: pip install python-docx docxtpl")
    sys.exit(1)


# ============ 数据类定义 ============

@dataclass
class PrintJob:
    """打印任务"""
    operator_id: int
    operator_name: str
    business_name: str
    document_path: str
    pdf_path: Optional[str] = None
    print_count: int = 0
    printed_at: Optional[str] = None
    status: str = "pending"  # pending, printed, failed


@dataclass
class GenerationResult:
    """生成结果"""
    success: bool
    document_path: Optional[str] = None
    pdf_path: Optional[str] = None
    print_job_id: Optional[str] = None
    errors: list = field(default_factory=list)
    warnings: list = field(default_factory=list)


# ============ 申请书生成器 ============

class ApplicationPrinter:
    """申请书生成和打印工具"""

    def __init__(
        self,
        template_path: Optional[str] = None,
        output_dir: str = "generated_applications"
    ):
        """初始化

        Args:
            template_path: Word模板路径
            output_dir: 输出目录
        """
        self.template_path = Path(template_path) if template_path else None
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # 打印记录
        self.print_log_path = self.output_dir / "print_log.json"
        self.print_log = self._load_print_log()

        logger.info(f"申请书生成器初始化: 输出目录={self.output_dir}")

    def _load_print_log(self) -> Dict:
        """加载打印记录"""
        if self.print_log_path.exists():
            try:
                with open(self.print_log_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"加载打印记录失败: {e}")
        return {}

    def _save_print_log(self):
        """保存打印记录"""
        try:
            with open(self.print_log_path, 'w', encoding='utf-8') as f:
                json.dump(self.print_log, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"保存打印记录失败: {e}")

    def generate_application(
        self,
        operator_data: Dict[str, Any],
        template_name: Optional[str] = None
    ) -> GenerationResult:
        """生成申请书

        Args:
            operator_data: 经营户数据
            template_name: 模板名称

        Returns:
            生成结果
        """
        result = GenerationResult(success=False)

        try:
            # 1. 数据验证
            if not self._validate_data(operator_data):
                result.errors.append("数据验证失败")
                return result

            # 2. 使用Jinja2模板生成（如果可用）
            # 或者使用 python-docx 直接生成
            doc_path = self._generate_word_document(operator_data, template_name)

            if not doc_path:
                result.errors.append("Word文档生成失败")
                return result

            result.document_path = str(doc_path)

            # 3. 转换为PDF（可选）
            pdf_path = self._convert_to_pdf(doc_path)
            if pdf_path:
                result.pdf_path = str(pdf_path)

            # 4. 创建打印任务
            print_job_id = self._create_print_job(operator_data, result.document_path, result.pdf_path)
            result.print_job_id = print_job_id

            result.success = True
            logger.info(f"申请书生成成功: {doc_path}")

        except Exception as e:
            result.errors.append(f"生成失败: {str(e)}")
            logger.error(f"申请书生成失败: {e}")

        return result

    def _validate_data(self, data: Dict) -> bool:
        """验证数据完整性"""
        required_fields = [
            'operator_name',
            'id_card',
            'business_name',
            'business_address'
        ]

        for field in required_fields:
            if not data.get(field):
                logger.warning(f"缺少必需字段: {field}")
                return False

        return True

    def _generate_word_document(
        self,
        data: Dict,
        template_name: Optional[str] = None
    ) -> Optional[Path]:
        """生成Word文档 - 使用官方模板

        Args:
            data: 经营户数据
            template_name: 模板名称

        Returns:
            文档路径
        """
        try:
            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            id_suffix = data.get('id_card', '')[-4:]
            filename = f"application_{id_suffix}_{timestamp}.docx"
            output_path = self.output_dir / filename

            # 使用官方模板
            template_path = Path("templates/个体工商户开业登记申请书.docx")

            if not template_path.exists():
                logger.warning(f"模板文件不存在: {template_path}, 使用默认生成方式")
                return self._generate_word_default(data, output_path)

            # 准备模板上下文数据
            context = self._prepare_template_context(data)

            # 加载模板并填充
            template = DocxTemplate(str(template_path))
            template.render(context)
            template.save(str(output_path))

            logger.info(f"Word文档生成(模板): {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Word文档生成失败: {e}")
            # 降级到默认生成方式
            return self._generate_word_default(data, output_path)

    def _generate_word_default(self, data: Dict, output_path: Path) -> Optional[Path]:
        """默认方式生成Word文档（当模板不可用时）

        Args:
            data: 经营户数据
            output_path: 输出路径

        Returns:
            文档路径
        """
        try:
            # 创建文档
            doc = Document()

            # 设置文档标题
            title = doc.add_paragraph()
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title.add_run('个体工商户登记（备案）申请书')
            run.bold = True
            run.font.size = Pt(16)

            # 添加内容
            self._add_application_content(doc, data)

            # 保存文档
            doc.save(str(output_path))

            logger.info(f"Word文档生成(默认): {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Word文档生成失败: {e}")
            return None

    def _prepare_template_context(self, data: Dict) -> Dict:
        """准备模板上下文数据

        Args:
            data: 原始经营户数据

        Returns:
            模板上下文字典
        """
        # 基本信息
        context = {
            'operator_name': data.get('operator_name', ''),
            'id_card': data.get('id_card', ''),
            'phone': data.get('phone', ''),
            'email': data.get('email', ''),
            'gender': data.get('gender', ''),
            'nation': data.get('nation', ''),
            'address': data.get('address', ''),

            # 经营信息
            'business_name': data.get('business_name', ''),
            'business_address': data.get('business_address', ''),
            'business_address_detail': data.get('business_address_detail', ''),
            'business_area': data.get('business_area', ''),
            'business_type': data.get('business_type', ''),
            'business_scope': data.get('business_scope', ''),
            'business_scope_licensed': data.get('business_scope_licensed', ''),
            'business_scope_general': data.get('business_scope_general', ''),

            # 从业和资金
            'employee_count': data.get('employee_count', 1),
            'registered_capital': data.get('registered_capital', '0.01'),

            # 场所信息
            'property_owner': data.get('property_owner', ''),
            'lease_start': data.get('lease_start', ''),
            'lease_end': data.get('lease_end', ''),
            'rent_amount': data.get('rent_amount', ''),

            # 日期
            'current_date': datetime.now().strftime('%Y年%m月%d日'),
        }

        return context

    def _add_application_content(self, doc: Document, data: Dict):
        """添加申请书内容

        Args:
            doc: Document对象
            data: 经营户数据
        """
        # 基本信息
        doc.add_paragraph('一、基本信息')
        doc.add_paragraph(f'经营者姓名：{data.get("operator_name", "")}')
        doc.add_paragraph(f'身份证号：{data.get("id_card", "")}')
        doc.add_paragraph(f'联系电话：{data.get("phone", "")}')
        doc.add_paragraph(f'电子邮箱：{data.get("email", "")}')
        doc.add_paragraph(f'住址：{data.get("address", "")}')

        # 经营信息
        doc.add_paragraph('\n二、经营信息')
        doc.add_paragraph(f'个体工商户名称：{data.get("business_name", "")}')
        doc.add_paragraph(f'经营场所地址：{data.get("business_address", "")}')

        # 详细地址
        if data.get("business_address_detail"):
            doc.add_paragraph(f'详细地址：{data["business_address_detail"]}')

        # 经营范围
        doc.add_paragraph('\n三、经营范围')
        doc.add_paragraph(f'经营范围：{data.get("business_scope", "")}')

        # 细化经营范围
        if data.get("business_scope_licensed"):
            doc.add_paragraph(f'许可经营项目：{data["business_scope_licensed"]}')
        if data.get("business_scope_general"):
            doc.add_paragraph(f'一般经营项目：{data["business_scope_general"]}')

        # 从业和资金
        doc.add_paragraph('\n四、从业情况')
        doc.add_paragraph(f'从业人数：{data.get("employee_count", "1")} 人')
        doc.add_paragraph(f'注册资金：{data.get("registered_capital", "0.01")} 万元')

        # 场所信息
        if data.get("property_owner"):
            doc.add_paragraph('\n五、场所信息')
            doc.add_paragraph(f'房产所有人/房东：{data["property_owner"]}')
            if data.get("lease_start"):
                doc.add_paragraph(f'租赁期限：{data["lease_start"]} 至 {data["lease_end"]}')
            if data.get("rent_amount"):
                doc.add_paragraph(f'租金：{data["rent_amount"]}')

        # 签名区
        doc.add_paragraph('\n' + '=' * 50)
        doc.add_paragraph('申请人（签字/盖章）：________________')
        doc.add_paragraph(f'日期：{datetime.now().strftime("%Y年%m月%d日")}')

    def _convert_to_pdf(self, doc_path: Path) -> Optional[Path]:
        """转换Word为PDF

        Args:
            doc_path: Word文档路径

        Returns:
            PDF路径
        """
        # 注意: 这需要额外的库支持
        # 例如: docx2pdf, win32com (Windows), LibreOffice

        # 简化版本：仅返回路径，实际转换需要额外配置
        pdf_path = doc_path.with_suffix('.pdf')

        logger.info(f"PDF转换暂未实现，路径: {pdf_path}")
        return None

    def _create_print_job(
        self,
        operator_data: Dict,
        document_path: str,
        pdf_path: Optional[str]
    ) -> str:
        """创建打印任务

        Args:
            operator_data: 经营户数据
            document_path: 文档路径
            pdf_path: PDF路径

        Returns:
            打印任务ID
        """
        import uuid

        job_id = str(uuid.uuid4())

        job = PrintJob(
            operator_id=operator_data.get('id', 0),
            operator_name=operator_data.get('operator_name', ''),
            business_name=operator_data.get('business_name', ''),
            document_path=document_path,
            pdf_path=pdf_path,
            status="pending"
        )

        self.print_log[job_id] = {
            "job_id": job_id,
            "operator_id": job.operator_id,
            "operator_name": job.operator_name,
            "business_name": job.business_name,
            "document_path": job.document_path,
            "pdf_path": job.pdf_path,
            "print_count": job.print_count,
            "printed_at": job.printed_at,
            "status": job.status,
            "created_at": datetime.now().isoformat()
        }

        self._save_print_log()

        return job_id

    def mark_printed(self, job_id: str, operator_id: int) -> bool:
        """标记为已打印

        Args:
            job_id: 打印任务ID
            operator_id: 经营户ID

        Returns:
            是否成功
        """
        try:
            # 更新打印记录
            if job_id in self.print_log:
                self.print_log[job_id]["status"] = "printed"
                self.print_log[job_id]["print_count"] += 1
                self.print_log[job_id]["printed_at"] = datetime.now().isoformat()

            self._save_print_log()

            logger.info(f"标记已打印: job_id={job_id}, operator_id={operator_id}")
            return True

        except Exception as e:
            logger.error(f"标记打印失败: {e}")
            return False

    def get_print_status(self, job_id: str) -> Optional[Dict]:
        """获取打印状态

        Args:
            job_id: 打印任务ID

        Returns:
            打印状态
        """
        return self.print_log.get(job_id)

    def list_print_jobs(
        self,
        operator_id: Optional[int] = None,
        status: Optional[str] = None
    ) -> list:
        """列出打印任务

        Args:
            operator_id: 经营户ID过滤
            status: 状态过滤

        Returns:
            打印任务列表
        """
        jobs = list(self.print_log.values())

        # 过滤
        if operator_id is not None:
            jobs = [j for j in jobs if j.get("operator_id") == operator_id]
        if status is not None:
            jobs = [j for j in jobs if j.get("status") == status]

        # 按创建时间排序
        jobs.sort(key=lambda x: x.get("created_at", ""), reverse=True)

        return jobs

    def prepare_for_print(self, job_id: str) -> Dict:
        """准备打印

        Args:
            job_id: 打印任务ID

        Returns:
            打印信息
        """
        job = self.print_log.get(job_id)

        if not job:
            return {"error": "打印任务不存在"}

        return {
            "job_id": job_id,
            "document_path": job.get("document_path"),
            "pdf_path": job.get("pdf_path"),
            "operator_name": job.get("operator_name"),
            "business_name": job.get("business_name"),
            "print_count": job.get("print_count", 0),
            "action": "open_document",  # 指示前端打开文档
            "message": f"请打开文档并打印，打印后请标记为已打印"
        }

    def export_print_record(self, output_path: str) -> bool:
        """导出打印记录

        Args:
            output_path: 输出文件路径

        Returns:
            是否成功
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(self.print_log, f, ensure_ascii=False, indent=2)

            logger.info(f"导出打印记录: {output_path}")
            return True

        except Exception as e:
            logger.error(f"导出打印记录失败: {e}")
            return False


# ============ 便捷函数 ============

def create_printer(
    template_path: Optional[str] = None,
    output_dir: str = "generated_applications"
) -> ApplicationPrinter:
    """创建打印工具实例

    Args:
        template_path: 模板路径
        output_dir: 输出目录

    Returns:
        ApplicationPrinter实例
    """
    return ApplicationPrinter(template_path, output_dir)


def quick_generate(operator_data: Dict) -> GenerationResult:
    """快速生成申请书

    Args:
        operator_data: 经营户数据

    Returns:
        生成结果
    """
    printer = ApplicationPrinter()
    return printer.generate_application(operator_data)


# ============ 测试代码 ============

if __name__ == "__main__":
    # 测试数据
    test_data = {
        "id": 1,
        "operator_name": "张三",
        "id_card": "450924199001011234",
        "phone": "13812345678",
        "email": "zhangsan@example.com",
        "address": "广西玉林市兴业县蒲塘镇XX村",
        "business_name": "玉林市兴业县张三便利店（个体工商户）",
        "business_address": "广西玉林市兴业县蒲塘镇XX街道XX号",
        "business_address_detail": "蒲塘镇开发区商业街第1间",
        "business_scope": "日用百货、烟酒零售",
        "business_scope_licensed": "小餐饮",
        "business_scope_general": "食品销售（仅销售预包装食品）",
        "employee_count": 2,
        "registered_capital": "5.0",
        "property_owner": "李四",
        "lease_start": "2024-01-01",
        "lease_end": "2029-12-31",
        "rent_amount": "2000元/月"
    }

    # 创建打印工具
    printer = ApplicationPrinter()

    # 生成申请书
    print("=" * 60)
    print("生成申请书测试")
    print("=" * 60)

    result = printer.generate_application(test_data)

    print(f"\n生成结果:")
    print(f"  成功: {result.success}")
    print(f"  文档路径: {result.document_path}")
    print(f"  PDF路径: {result.pdf_path}")
    print(f"  打印任务ID: {result.print_job_id}")

    if result.errors:
        print(f"  错误: {', '.join(result.errors)}")

    # 准备打印
    if result.print_job_id:
        print(f"\n打印信息:")
        print_info = printer.prepare_for_print(result.print_job_id)
        for key, value in print_info.items():
            print(f"  {key}: {value}")
