#!/usr/bin/env python3
"""
智能填充李奕凤版模板 v3 - 最终版
- 红色字体 RGB(255,0,0) = 示例数据（需要替换）
- 黑色/默认字体 = 字段名（保留）
"""

import json
import datetime
from pathlib import Path
from docx import Document

def is_red_font(run):
    """判断是否为红色字体 RGB(255,0,0)"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        return r == 255 and g == 0 and b == 0
    return False

def get_cell_black_text(cell):
    """获取单元格中的黑色/默认文字（字段名）"""
    black_text = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if not is_red_font(run):  # 非红色字体
                text = run.text.strip()
                if text:
                    black_text.append(text)
    return " ".join(black_text)

def fill_liyifeng_template_final(data, template_file="（李奕凤）个体工商户开业登记申请书（模板）.docx"):
    """智能填充李奕凤版模板"""
    try:
        print(f"正在读取模板: {template_file}")
        doc = Document(template_file)

        print("正在智能填充红色字体字段...")

        # 定义数据映射（根据黑色字体上下文匹配）
        data_mapping = {
            "个体工商户名称": "business_name",
            "经营者": "operator_name",
            "联系电话": "phone",
            "联系": "phone",
            "电话": "phone",
            "电子邮箱": "email",
            "邮箱": "email",
            "经营场所": "business_address",
            "住所": "business_address",
            "地址": "business_address",
            "邮政编码": "postal_code",
            "邮编": "postal_code",
            "从业人数": "employee_count",
            "从业": "employee_count",
            "人数": "employee_count",
            "注册资金": "registered_capital",
            "注册": "registered_capital",
            "资金": "registered_capital",
            "身份证": "id_card",
            "身份证号码": "id_card",
            "性别": "gender",
            "民族": "nation",
            "政治面貌": "political_status",
            "政治": "political_status",
            "文化程度": "education",
            "文化": "education",
            "学历": "education",
            "经营范围": "business_scope",
            "经营": "business_scope",
            "范围": "business_scope",
            "经营期限": "operation_period",
            "期限": "operation_period"
        }

        filled_count = 0
        filled_details = []

        # 遍历所有表格
        for table_idx, table in enumerate(doc.tables):
            print(f"\n处理表格 {table_idx}...")

            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # 获取单元格中的黑色文字（上下文字段名）
                    black_text = get_cell_black_text(cell)

                    # 查找匹配的数据字段
                    matched_key = None
                    for keyword, data_key in data_mapping.items():
                        if keyword in black_text:
                            matched_key = data_key
                            break

                    # 如果找到匹配，替换红色字体
                    if matched_key and matched_key in data:
                        value = str(data.get(matched_key, ""))
                        clean_value = ''.join(c for c in value if ord(c) < 0xD800 or ord(c) > 0xDFFF)

                        # 替换所有红色字体的run
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if is_red_font(run):
                                    old_text = run.text
                                    if old_text.strip():
                                        run.text = clean_value
                                        filled_count += 1
                                        filled_details.append(f"  [{matched_key}] '{old_text[:20]}...' -> '{clean_value}'")
                                        print(f"  ✓ {matched_key}: {old_text[:20]}... -> {clean_value}")

        print(f"\n完成统计:")
        print(f"  总共填充: {filled_count} 处红色字段")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_李奕凤版_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")
        print(f"[说明] 李奕凤版模板，红色字体已替换为实际数据")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] {data_file}")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数"""
    # 测试数据
    test_data = {
        'business_name': '钱十一餐厅',
        'operator_name': '钱十一',
        'phone': '13222222222',
        'email': 'qianshiyi@test.com',
        'business_address': '西安市雁塔区小寨西路168号',
        'postal_code': '710000',
        'employee_count': '15',
        'registered_capital': '600000',
        'id_card': '610101198404043333',
        'gender': '男',
        'nation': '汉族',
        'political_status': '党员',
        'education': '本科',
        'business_scope': '中餐服务；酒水饮料',
        'operation_period': '长期'
    }

    print("=" * 70)
    print("李奕凤版模板填充工具 v3 - 最终版")
    print("红色字体 RGB(255,0,0) = 示例数据（自动替换）")
    print("黑色字体 = 字段名（保留不变）")
    print("=" * 70)

    print("\n测试数据：")
    for key, value in test_data.items():
        print(f"  {key}: {value}")

    result = fill_liyifeng_template_final(test_data)

    if result:
        print("\n" + "=" * 70)
        print("测试成功！")
        print(f"生成的文件: {result}")
        print("=" * 70)

        # 用WPS打开
        print("\n正在用WPS打开文档...")
        import subprocess
        subprocess.Popen(['cmd', '/c', 'start', '', str(result)], shell=True)

if __name__ == "__main__":
    main()
