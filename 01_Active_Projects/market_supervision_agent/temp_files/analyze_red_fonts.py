#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析模板中的所有红色字体内容
"""

from docx import Document
import json

def is_red_font(run):
    """判断是否为红色字体"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        return r == 255 and g == 0 and b == 0
    return False

def analyze_red_fonts():
    """详细分析红色字体"""

    print("=" * 70)
    print("红色字体详细分析")
    print("=" * 70)

    # 读取原模板
    src = Document(r'c:\Users\flyskyson\Desktop\（李奕凤）个体工商户开业登记申请书（模板-待修改）.docx')

    red_font_items = []

    # 分析段落
    for para_idx, para in enumerate(src.paragraphs):
        for run_idx, run in enumerate(para.runs):
            if is_red_font(run):
                red_font_items.append({
                    'location': f'Paragraph {para_idx}, Run {run_idx}',
                    'text': run.text,
                    'type': 'paragraph'
                })

    # 分析表格
    for table_idx, table in enumerate(src.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(para.runs):
                        if is_red_font(run):
                            red_font_items.append({
                                'location': f'Table{table_idx} Row{row_idx} Cell{cell_idx}',
                                'text': run.text,
                                'type': 'table'
                            })

    # 去重
    seen = set()
    unique_items = []
    for item in red_font_items:
        text = item['text']
        if text and text not in seen:
            seen.add(text)
            unique_items.append(item)

    print(f"\n找到 {len(unique_items)} 个不同的红色字体内容:\n")

    for idx, item in enumerate(unique_items, 1):
        text_preview = item['text'][:50] if len(item['text']) > 50 else item['text']
        print(f"{idx}. [{item['location']}]")
        print(f"   内容: {text_preview}")
        print()

    # 保存分析结果
    output_file = r'c:\Users\flyskyson\Desktop\red_fonts_analysis.json'
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump({
            'total': len(unique_items),
            'items': unique_items
        }, f, ensure_ascii=False, indent=2)

    print(f"分析结果已保存到: {output_file}")
    print("=" * 70)

if __name__ == "__main__":
    analyze_red_fonts()
