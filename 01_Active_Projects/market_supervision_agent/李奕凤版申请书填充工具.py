#!/usr/bin/env python3
"""
李奕凤版个体工商户开业登记申请书填充工具 - 生产版 v2
- 自动删除红色字体（示例数据）
- 智能识别字段名并填充实际数据
- 完美保持模板格式
- 自动用WPS打开生成的文档
- 支持命令行参数和批量处理
"""

import json
import datetime
import sys
import argparse
from pathlib import Path
from docx import Document

# ============ 核心功能 ============

def is_red_font(run):
    """判断是否为红色字体 RGB(255,0,0)"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        return r == 255 and g == 0 and b == 0
    return False

def cell_has_red_font(cell):
    """检查单元格是否有红色字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if is_red_font(run):
                return True
    return False

def get_cell_text_clean(cell):
    """获取单元格的纯文本（用于字段名识别）"""
    text_parts = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if not is_red_font(run) and run.text.strip():
                text_parts.append(run.text.strip())
    return "".join(text_parts)

def fill_template(data, template_file="（李奕凤）个体工商户开业登记申请书（模板）.docx",
                  output_dir="output", auto_open=True):
    """
    填充李奕凤版模板

    参数:
        data: 包含申请信息的字典
        template_file: 模板文件路径
        output_dir: 输出目录
        auto_open: 是否自动打开文档

    返回:
        生成的文件路径，失败返回None
    """
    try:
        print(f"正在读取模板: {template_file}")
        doc = Document(template_file)

        print("正在智能填充字段...")

        # 字段识别规则
        field_patterns = [
            (["个体工商户名称", "名称", "字号"], "business_name"),
            (["经营者姓名", "经营者"], "operator_name"),
            (["性别"], "gender"),
            (["民族"], "nation"),
            (["政治面貌", "政治"], "political_status"),
            (["文化程度", "学历", "教育"], "education"),
            (["身份证号码", "身份证"], "id_card"),
            (["联系电话", "电话"], "phone"),
            (["电子邮箱", "邮箱"], "email"),
            (["邮政编码", "邮编"], "postal_code"),
            (["经营场所", "住所", "地址"], "business_address"),
            (["从业人数", "人数"], "employee_count"),
            (["注册资金", "资金"], "registered_capital"),
            (["经营范围"], "business_scope"),
            (["经营期限", "期限"], "operation_period"),
        ]

        deleted_count = 0
        filled_count = 0
        filled_fields = set()

        # 遍历所有表格
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if not cell_has_red_font(cell):
                        continue

                    # 查找字段名
                    field_name_text = ""

                    # 优先左侧单元格
                    if cell_idx > 0:
                        left_cell = row.cells[cell_idx - 1]
                        field_name_text = get_cell_text_clean(left_cell)

                    # 删除红色字体
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if is_red_font(run):
                                run.text = ""
                                deleted_count += 1

                    # 匹配并填充
                    matched_key = None
                    for field_names, data_key in field_patterns:
                        if data_key not in filled_fields:
                            for field_name in field_names:
                                if field_name in field_name_text:
                                    matched_key = data_key
                                    filled_fields.add(data_key)
                                    break
                            if matched_key:
                                break

                    if matched_key and matched_key in data:
                        value = str(data[matched_key])
                        if value and cell.paragraphs:
                            para = cell.paragraphs[0]
                            para.text = ""
                            para.clear()

                            new_run = para.add_run(value)
                            new_run.font.name = '宋体'
                            new_run.font.size = 209712

                            filled_count += 1
                            print(f"  [OK] {matched_key}: {value}")

        # 保存文件
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        output_file = output_path / f"{business_name}_开业登记_{timestamp}.docx"

        doc.save(str(output_file))

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")
        print(f"[数据] {data_file}")
        print(f"[统计] 删除红色字体{deleted_count}处，填充数据{filled_count}处")

        # 自动打开
        if auto_open:
            import subprocess
            subprocess.Popen(['cmd', '/c', 'start', '', str(output_file)], shell=True)
            print("\n正在用WPS打开文档...")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

# ============ 批量处理 ============

def batch_process(data_file, template_file=None, output_dir="output", auto_open=True):
    """批量处理JSON文件中的多条数据"""
    try:
        with open(data_file, 'r', encoding='utf-8') as f:
            data_list = json.load(f)

        # 如果是单个字典，转换为列表
        if isinstance(data_list, dict):
            data_list = [data_list]

        print(f"\n[批量处理] 共 {len(data_list)} 条数据\n")

        results = []
        for idx, data in enumerate(data_list, 1):
            print(f"\n处理第 {idx}/{len(data_list)} 条:")
            result = fill_template(data, template_file, output_dir, auto_open=False)
            if result:
                results.append(result)

        # 全部处理完后，打开最后一个文件
        if auto_open and results:
            import subprocess
            subprocess.Popen(['cmd', '/c', 'start', '', str(results[-1])], shell=True)

        print(f"\n[完成] 共生成 {len(results)} 个文档")
        return results

    except Exception as e:
        print(f"[错误] 批量处理失败: {e}")
        import traceback
        traceback.print_exc()
        return []

# ============ 主程序 ============

def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='李奕凤版个体工商户开业登记申请书填充工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  # 使用测试数据
  python 李奕凤版申请书填充工具.py --test

  # 从JSON文件加载数据
  python 李奕凤版申请书填充工具.py --data test_data.json

  # 批量处理
  python 李奕凤版申请书填充工具.py --batch data_list.json

  # 指定模板和输出目录
  python 李奕凤版申请书填充工具.py --data test.json --template mytemplate.docx --output myoutput
        """
    )

    parser.add_argument('--data', type=str, help='JSON数据文件路径')
    parser.add_argument('--batch', type=str, help='批量处理的JSON文件路径')
    parser.add_argument('--template', type=str, default='（李奕凤）个体工商户开业登记申请书（模板）.docx',
                        help='模板文件路径')
    parser.add_argument('--output', type=str, default='output', help='输出目录')
    parser.add_argument('--test', action='store_true', help='使用测试数据')
    parser.add_argument('--no-open', action='store_true', help='不自动打开文档')

    args = parser.parse_args()

    print("=" * 70)
    print("李奕凤版个体工商户开业登记申请书填充工具")
    print("生产版 - 自动填充，保持格式")
    print("=" * 70)

    # 检查模板文件
    if not Path(args.template).exists():
        print(f"\n[错误] 模板文件不存在: {args.template}")
        print("\n提示: 请将模板文件放在当前目录，或使用 --template 参数指定")
        return 1

    auto_open = not args.no_open

    # 根据参数执行不同操作
    if args.test:
        # 使用测试数据
        test_data = {
            'business_name': '赵六便利店',
            'operator_name': '赵六',
            'phone': '13444444444',
            'email': 'zhaoliu@test.com',
            'business_address': '深圳市福田区深南大道666号',
            'postal_code': '518000',
            'employee_count': '4',
            'registered_capital': '60000',
            'id_card': '440101198704044444',
            'gender': '男',
            'nation': '汉族',
            'political_status': '群众',
            'education': '大专',
            'business_scope': '便利店零售；日用百货',
            'operation_period': '长期'
        }

        print("\n[测试模式] 使用预设测试数据\n")
        result = fill_template(test_data, args.template, args.output, auto_open)

    elif args.batch:
        # 批量处理
        results = batch_process(args.batch, args.template, args.output, auto_open)
        if not results:
            return 1

    elif args.data:
        # 从文件加载
        try:
            with open(args.data, 'r', encoding='utf-8') as f:
                data = json.load(f)

            print(f"\n[加载数据] 从文件: {args.data}\n")
            result = fill_template(data, args.template, args.output, auto_open)

        except Exception as e:
            print(f"\n[错误] 加载数据失败: {e}")
            return 1

    else:
        # 显示帮助
        parser.print_help()
        print("\n快速开始:")
        print("  python 李奕芙版申请书填充工具.py --test")
        return 0

    print("\n" + "=" * 70)
    print("处理完成！")
    print("=" * 70)
    return 0

if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        print("\n\n[中断] 用户取消操作")
        sys.exit(0)
    except Exception as e:
        print(f"\n[异常] {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
