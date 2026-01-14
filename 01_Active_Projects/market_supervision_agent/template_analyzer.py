#!/usr/bin/env python3
"""
Wordæ¨¡æ¿åˆ†æå·¥å…·
ç”¨äºåˆ†æä¸ªä½“å·¥å•†æˆ·ç™»è®°ç”³è¯·ä¹¦çš„Wordæ¨¡æ¿ç»“æ„
"""

import sys
from pathlib import Path
from typing import Dict, List, Tuple, Set
import json

try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table, _Cell
except ImportError:
    print("âŒ éœ€è¦å®‰è£… python-docx åº“")
    print("è¯·è¿è¡Œ: pip install python-docx")
    sys.exit(1)

class TemplateAnalyzer:
    """Wordæ¨¡æ¿åˆ†æå™¨"""

    def __init__(self):
        self.placeholder_patterns = [
            "{{", "}}",  # åŒèŠ±æ‹¬å·
            "[", "]",    # æ–¹æ‹¬å·
            "<<", ">>",  # åŒå°–æ‹¬å·
            "__", "__",  # åŒä¸‹åˆ’çº¿
            "##", "##",  # åŒäº•å·
        ]

    def analyze_template(self, template_path: str) -> Dict:
        """
        åˆ†æWordæ¨¡æ¿ç»“æ„

        Args:
            template_path: Wordæ¨¡æ¿æ–‡ä»¶è·¯å¾„

        Returns:
            åŒ…å«æ¨¡æ¿åˆ†æç»“æœçš„å­—å…¸
        """
        template_path = Path(template_path)
        if not template_path.exists():
            print(f"æ¨¡æ¿æ–‡ä»¶ä¸å­˜åœ¨: {template_path}")
            return {}

        print(f"\næ­£åœ¨åˆ†ææ¨¡æ¿: {template_path.name}")
        print("=" * 60)

        try:
            doc = Document(str(template_path))
            analysis = {
                "template_name": template_path.name,
                "template_size_kb": template_path.stat().st_size / 1024,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "sections_count": len(doc.sections),
                "placeholders_found": [],
                "suggested_fields": [],
                "template_structure": [],
                "extracted_text": []
            }

            # åˆ†ææ®µè½
            print(f"\næ®µè½åˆ†æ:")
            print(f"  æ®µè½æ€»æ•°: {len(doc.paragraphs)}")

            for i, paragraph in enumerate(doc.paragraphs[:10]):  # åªåˆ†æå‰10ä¸ªæ®µè½
                if paragraph.text.strip():
                    analysis["extracted_text"].append({
                        "type": "paragraph",
                        "index": i,
                        "text": paragraph.text.strip(),
                        "style": paragraph.style.name if paragraph.style else "Normal"
                    })

                    # æ£€æŸ¥å ä½ç¬¦
                    placeholders = self._find_placeholders(paragraph.text)
                    if placeholders:
                        analysis["placeholders_found"].extend(placeholders)
                        print(f"  æ®µè½ {i+1}: æ‰¾åˆ°å ä½ç¬¦ {placeholders}")

            # åˆ†æè¡¨æ ¼
            print(f"\nè¡¨æ ¼åˆ†æ:")
            print(f"  è¡¨æ ¼æ€»æ•°: {len(doc.tables)}")

            for i, table in enumerate(doc.tables):
                table_info = {
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "cells": []
                }

                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            table_info["cells"].append({
                                "row": row_idx,
                                "col": col_idx,
                                "text": cell_text
                            })

                            # æ£€æŸ¥å ä½ç¬¦
                            placeholders = self._find_placeholders(cell_text)
                            if placeholders:
                                analysis["placeholders_found"].extend(placeholders)
                                print(f"  è¡¨æ ¼ {i+1}, å•å…ƒæ ¼({row_idx+1},{col_idx+1}): æ‰¾åˆ°å ä½ç¬¦ {placeholders}")

                analysis["template_structure"].append({
                    "type": "table",
                    "info": table_info
                })

            # å»é‡å ä½ç¬¦
            analysis["placeholders_found"] = list(set(analysis["placeholders_found"]))

            # æ ¹æ®å ä½ç¬¦å»ºè®®å­—æ®µ
            analysis["suggested_fields"] = self._suggest_fields(analysis["placeholders_found"])

            # æ ¹æ®æ–‡æœ¬å†…å®¹å»ºè®®å­—æ®µ
            text_based_fields = self._suggest_fields_from_text(analysis["extracted_text"])
            analysis["suggested_fields"].extend([f for f in text_based_fields if f not in analysis["suggested_fields"]])

            print(f"\nåˆ†æå®Œæˆï¼")
            print(f"  æ‰¾åˆ°å ä½ç¬¦: {len(analysis['placeholders_found'])} ä¸ª")
            print(f"  å»ºè®®å­—æ®µ: {len(analysis['suggested_fields'])} ä¸ª")

            return analysis

        except Exception as e:
            print(f"åˆ†ææ¨¡æ¿å¤±è´¥: {str(e)}")
            return {}

    def _find_placeholders(self, text: str) -> List[str]:
        """åœ¨æ–‡æœ¬ä¸­æŸ¥æ‰¾å ä½ç¬¦"""
        placeholders = []

        # æ£€æŸ¥å„ç§å ä½ç¬¦æ ¼å¼
        for i in range(0, len(self.placeholder_patterns), 2):
            start_pattern = self.placeholder_patterns[i]
            end_pattern = self.placeholder_patterns[i + 1]

            start_idx = text.find(start_pattern)
            while start_idx != -1:
                end_idx = text.find(end_pattern, start_idx + len(start_pattern))
                if end_idx != -1:
                    placeholder = text[start_idx + len(start_pattern):end_idx].strip()
                    if placeholder:
                        placeholders.append(placeholder)
                    start_idx = text.find(start_pattern, end_idx + len(end_pattern))
                else:
                    break

        return placeholders

    def _suggest_fields(self, placeholders: List[str]) -> List[Dict]:
        """æ ¹æ®å ä½ç¬¦å»ºè®®å­—æ®µ"""
        field_mapping = {
            # åç§°ç›¸å…³
            "name": ["åç§°", "å§“å", "name", "å§“å", "ç»è¥è€…å§“å", "ä¸ªä½“å·¥å•†æˆ·åç§°"],
            "business_name": ["ä¸ªä½“å·¥å•†æˆ·åç§°", "ä¼ä¸šåç§°", "åº—é“ºåç§°", "å­—å·", "å•†å·"],
            "operator_name": ["ç»è¥è€…å§“å", "è´Ÿè´£äºº", "æ³•äºº", "ä¸šä¸»", "ç»è¥è€…"],

            # è¯ä»¶ç›¸å…³
            "id_card": ["èº«ä»½è¯", "è¯ä»¶å·ç ", "èº«ä»½è¯å·", "èº«ä»½è¯å·ç ", "è¯ä»¶å·"],
            "credit_code": ["ç»Ÿä¸€ç¤¾ä¼šä¿¡ç”¨ä»£ç ", "ä¿¡ç”¨ä»£ç ", "ç¤¾ä¼šä¿¡ç”¨ä»£ç ", "ç»„ç»‡æœºæ„ä»£ç "],

            # è”ç³»ä¿¡æ¯
            "phone": ["ç”µè¯", "è”ç³»ç”µè¯", "æ‰‹æœº", "æ‰‹æœºå·", "è”ç³»æ–¹å¼"],
            "email": ["é‚®ç®±", "ç”µå­é‚®ç®±", "email", "ç”µå­é‚®ä»¶"],
            "address": ["åœ°å€", "ç»è¥åœºæ‰€", "ä½æ‰€", "è¥ä¸šåœ°å€", "åœºæ‰€"],

            # ç»è¥ä¿¡æ¯
            "business_scope": ["ç»è¥èŒƒå›´", "ç»è¥é¡¹ç›®", "ä¸šåŠ¡èŒƒå›´", "ä¸»è¥é¡¹ç›®"],
            "registered_capital": ["èµ„é‡‘", "æ³¨å†Œèµ„æœ¬", "å‡ºèµ„é¢", "èµ„é‡‘æ•°é¢", "æŠ•èµ„é¢"],
            "business_type": ["è¡Œä¸š", "è¡Œä¸šç±»å‹", "ç»è¥ç±»å‹", "ä¼ä¸šç±»å‹"],

            # æ—¶é—´ç›¸å…³
            "application_date": ["æ—¥æœŸ", "ç”³è¯·æ—¥æœŸ", "æäº¤æ—¥æœŸ", "åŠç†æ—¥æœŸ"],
            "start_date": ["å¼€ä¸šæ—¥æœŸ", "æˆç«‹æ—¥æœŸ", "èµ·å§‹æ—¥æœŸ"],
            "end_date": ["æˆªæ­¢æ—¥æœŸ", "ç»“æŸæ—¥æœŸ", "æœ‰æ•ˆæœŸè‡³"],
        }

        suggested_fields = []
        used_fields = set()

        for placeholder in placeholders:
            placeholder_lower = placeholder.lower()

            # ç›´æ¥åŒ¹é…
            for field_name, keywords in field_mapping.items():
                if field_name in used_fields:
                    continue

                for keyword in keywords:
                    if keyword in placeholder or keyword in placeholder_lower:
                        suggested_fields.append({
                            "field_name": field_name,
                            "display_name": keywords[0],
                            "placeholder": placeholder,
                            "description": f"å¯¹åº”å ä½ç¬¦: {placeholder}",
                            "required": True
                        })
                        used_fields.add(field_name)
                        break

        # æ·»åŠ æ ‡å‡†å­—æ®µï¼ˆå¦‚æœè¿˜æ²¡æ·»åŠ ï¼‰
        standard_fields = [
            {
                "field_name": "business_name",
                "display_name": "ä¸ªä½“å·¥å•†æˆ·åç§°",
                "placeholder": "{{ä¸ªä½“å·¥å•†æˆ·åç§°}}",
                "description": "ä¸ªä½“å·¥å•†æˆ·çš„å…¨ç§°",
                "required": True
            },
            {
                "field_name": "operator_name",
                "display_name": "ç»è¥è€…å§“å",
                "placeholder": "{{ç»è¥è€…å§“å}}",
                "description": "ç»è¥è€…çš„å§“å",
                "required": True
            },
            {
                "field_name": "id_card",
                "display_name": "èº«ä»½è¯å·ç ",
                "placeholder": "{{èº«ä»½è¯å·ç }}",
                "description": "ç»è¥è€…çš„èº«ä»½è¯å·ç ",
                "required": True
            },
            {
                "field_name": "business_address",
                "display_name": "ç»è¥åœºæ‰€",
                "placeholder": "{{ç»è¥åœºæ‰€}}",
                "description": "ä¸ªä½“å·¥å•†æˆ·çš„ç»è¥åœ°å€",
                "required": True
            },
            {
                "field_name": "phone",
                "display_name": "è”ç³»ç”µè¯",
                "placeholder": "{{è”ç³»ç”µè¯}}",
                "description": "ç»è¥è€…çš„è”ç³»ç”µè¯",
                "required": True
            },
            {
                "field_name": "business_scope",
                "display_name": "ç»è¥èŒƒå›´",
                "placeholder": "{{ç»è¥èŒƒå›´}}",
                "description": "ä¸ªä½“å·¥å•†æˆ·çš„ç»è¥èŒƒå›´",
                "required": True
            },
            {
                "field_name": "registered_capital",
                "display_name": "èµ„é‡‘æ•°é¢",
                "placeholder": "{{èµ„é‡‘æ•°é¢}}",
                "description": "ä¸ªä½“å·¥å•†æˆ·çš„èµ„é‡‘æ•°é¢ï¼ˆå•ä½ï¼šå…ƒï¼‰",
                "required": True
            },
            {
                "field_name": "application_date",
                "display_name": "ç”³è¯·æ—¥æœŸ",
                "placeholder": "{{ç”³è¯·æ—¥æœŸ}}",
                "description": "ç”³è¯·æ—¥æœŸï¼ˆè‡ªåŠ¨å¡«å……ï¼‰",
                "required": False
            }
        ]

        for field in standard_fields:
            if field["field_name"] not in used_fields:
                suggested_fields.append(field)
                used_fields.add(field["field_name"])

        return suggested_fields

    def _suggest_fields_from_text(self, extracted_text: List[Dict]) -> List[Dict]:
        """ä»æå–çš„æ–‡æœ¬ä¸­å»ºè®®å­—æ®µ"""
        suggested_fields = []

        # å¸¸è§çš„å…³é”®è¯æ¨¡å¼
        keyword_patterns = {
            "ä¸ªä½“å·¥å•†æˆ·åç§°": ["åç§°ï¼š", "åç§°ä¸º", "åç§°:", "ä¸ªä½“å·¥å•†æˆ·"],
            "ç»è¥è€…å§“å": ["ç»è¥è€…ï¼š", "ç»è¥è€…å§“åï¼š", "è´Ÿè´£äººï¼š", "æ³•äººï¼š"],
            "èº«ä»½è¯å·ç ": ["èº«ä»½è¯ï¼š", "èº«ä»½è¯å·ï¼š", "è¯ä»¶å·ç ï¼š", "èº«ä»½è¯å·ç ï¼š"],
            "ç»è¥åœºæ‰€": ["åœºæ‰€ï¼š", "åœ°å€ï¼š", "ç»è¥åœºæ‰€ï¼š", "ä½æ‰€ï¼š"],
            "ç»è¥èŒƒå›´": ["èŒƒå›´ï¼š", "ç»è¥èŒƒå›´ï¼š", "ç»è¥é¡¹ç›®ï¼š", "ä¸šåŠ¡èŒƒå›´ï¼š"],
            "èµ„é‡‘æ•°é¢": ["èµ„é‡‘ï¼š", "æ³¨å†Œèµ„æœ¬ï¼š", "å‡ºèµ„é¢ï¼š", "èµ„é‡‘æ•°é¢ï¼š"],
            "è”ç³»ç”µè¯": ["ç”µè¯ï¼š", "è”ç³»ç”µè¯ï¼š", "æ‰‹æœºï¼š", "è”ç³»æ–¹å¼ï¼š"]
        }

        for item in extracted_text:
            text = item["text"]
            for field_name, patterns in keyword_patterns.items():
                for pattern in patterns:
                    if pattern in text:
                        # æå–å­—æ®µå€¼ï¼ˆå¦‚æœå­˜åœ¨ï¼‰
                        value_start = text.find(pattern) + len(pattern)
                        value_end = text.find("ã€‚", value_start) if "ã€‚" in text[value_start:] else len(text)
                        example_value = text[value_start:value_end].strip()

                        suggested_fields.append({
                            "field_name": field_name.lower().replace("ï¼š", "").replace(":", ""),
                            "display_name": field_name,
                            "placeholder": f"{{{{{field_name}}}}}",
                            "description": f"ä»æ¨¡æ¿ä¸­è¯†åˆ«ï¼Œç¤ºä¾‹å€¼: {example_value[:20]}..." if example_value else f"ä»æ¨¡æ¿ä¸­è¯†åˆ«",
                            "required": True
                        })
                        break

        return suggested_fields

    def generate_filling_code(self, analysis: Dict, output_file: str = "template_filler.py"):
        """ç”Ÿæˆæ¨¡æ¿å¡«å……ä»£ç """
        if not analysis:
            print("âŒ æ²¡æœ‰åˆ†æç»“æœï¼Œæ— æ³•ç”Ÿæˆä»£ç ")
            return

        template_name = analysis.get("template_name", "unknown_template.docx")
        suggested_fields = analysis.get("suggested_fields", [])

        code = f'''#!/usr/bin/env python3
"""
Wordæ¨¡æ¿è‡ªåŠ¨å¡«å……å·¥å…·
ä¸“ä¸ºæ¨¡æ¿: {template_name}
è‡ªåŠ¨ç”Ÿæˆæ—¥æœŸ: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
"""

import sys
from pathlib import Path
import json
import datetime

try:
    from docx import Document
except ImportError:
    print("âŒ éœ€è¦å®‰è£… python-docx åº“")
    print("è¯·è¿è¡Œ: pip install python-docx")
    sys.exit(1)

class TemplateFiller:
    """{template_name} æ¨¡æ¿å¡«å……å™¨"""

    def __init__(self, template_path: str = "{template_name}"):
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            print(f"âŒ æ¨¡æ¿æ–‡ä»¶ä¸å­˜åœ¨: {{template_path}}")
            sys.exit(1)

    def fill_template(self, data: dict, output_path: str = None):
        """
        å¡«å……æ¨¡æ¿

        Args:
            data: åŒ…å«å­—æ®µæ•°æ®çš„å­—å…¸
            output_path: è¾“å‡ºæ–‡ä»¶è·¯å¾„ï¼Œå¦‚æœä¸ºNoneåˆ™è‡ªåŠ¨ç”Ÿæˆ
        """
        try:
            # è¯»å–æ¨¡æ¿
            doc = Document(str(self.template_path))

            # å‡†å¤‡æ›¿æ¢æ˜ å°„
            replacements = {{
'''

        # æ·»åŠ å­—æ®µæ˜ å°„
        for field in suggested_fields:
            field_name = field["field_name"]
            placeholder = field["placeholder"]
            code += f'                "{placeholder}": data.get("{field_name}", ""),\n'

        code += '''            }}

            # æ·»åŠ è‡ªåŠ¨ç”Ÿæˆçš„å­—æ®µ
            if "application_date" not in data:
                replacements["{{ç”³è¯·æ—¥æœŸ}}"] = datetime.datetime.now().strftime("%Yå¹´%mæœˆ%dæ—¥")

            # åœ¨æ®µè½ä¸­æ›¿æ¢
            for paragraph in doc.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, new_text)

            # åœ¨è¡¨æ ¼ä¸­æ›¿æ¢
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for old_text, new_text in replacements.items():
                            if old_text in cell.text:
                                cell.text = cell.text.replace(old_text, new_text)

            # ç”Ÿæˆè¾“å‡ºæ–‡ä»¶å
            if output_path is None:
                business_name = data.get("business_name", "æœªçŸ¥")
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                output_dir = Path("filled_templates")
                output_dir.mkdir(exist_ok=True)
                output_path = output_dir / f"{business_name}_{timestamp}.docx"

            # ä¿å­˜æ–‡ä»¶
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

            print(f"âœ… æ¨¡æ¿å¡«å……æˆåŠŸï¼")
            print(f"ğŸ“ æ–‡ä»¶ä¿å­˜ä½ç½®: {{output_path}}")

            return output_path

        except Exception as e:
            print(f"âŒ å¡«å……æ¨¡æ¿å¤±è´¥: {{str(e)}}")
            return None

    def get_required_fields(self):
        """è·å–å¿…å¡«å­—æ®µåˆ—è¡¨"""
        return [
'''

        # æ·»åŠ å¿…å¡«å­—æ®µè¯´æ˜
        for field in suggested_fields:
            if field.get("required", True):
                code += f'            "{field["field_name"]}",  # {field["display_name"]}\n'

        code += '''        ]

    def get_all_fields(self):
        """è·å–æ‰€æœ‰å­—æ®µåˆ—è¡¨"""
        return [
'''

        # æ·»åŠ æ‰€æœ‰å­—æ®µè¯´æ˜
        for field in suggested_fields:
            code += f'            {{"name": "{field["field_name"]}", "display": "{field["display_name"]}", "required": {field.get("required", True)}}},\n'

        code += '''        ]

def collect_data_interactive():
    """äº¤äº’å¼æ”¶é›†æ•°æ®"""
    print("ğŸ“ è¯·å¡«å†™ä»¥ä¸‹ä¿¡æ¯ï¼ˆæŒ‰Enterè·³è¿‡å¯é€‰å­—æ®µï¼‰ï¼š")
    data = {}

'''

        # æ·»åŠ äº¤äº’å¼æ•°æ®æ”¶é›†ä»£ç 
        for field in suggested_fields:
            field_name = field["field_name"]
            display_name = field["display_name"]
            required = field.get("required", True)

            if required:
                code += f'    data["{field_name}"] = input("{display_name}: ").strip()\n'
                code += f'    while not data["{field_name}"]:\n'
                code += f'        print("âš ï¸  {display_name}æ˜¯å¿…å¡«é¡¹")\n'
                code += f'        data["{field_name}"] = input("{display_name}: ").strip()\n'
            else:
                code += f'    value = input("{display_name}ï¼ˆå¯é€‰ï¼‰: ").strip()\n'
                code += f'    if value:\n'
                code += f'        data["{field_name}"] = value\n'
            code += '\n'

        code += '''    return data

def main():
    """ä¸»å‡½æ•°"""
    print("=" * 60)
    print(f"  {template_name} æ¨¡æ¿å¡«å……å·¥å…·")
    print("=" * 60)

    filler = TemplateFiller()

    print("\\nğŸ“‹ æ¨¡æ¿å­—æ®µè¯´æ˜ï¼š")
    fields = filler.get_all_fields()
    for field in fields:
        required = "å¿…å¡«" if field["required"] else "å¯é€‰"
        print(f"  â€¢ {field['display']} ({required})")

    print("\\n1. äº¤äº’å¼å¡«å†™")
    print("2. ä»JSONæ–‡ä»¶åŠ è½½")
    print("3. é€€å‡º")

    choice = input("\\nè¯·é€‰æ‹©æ“ä½œ (1-3): ").strip()

    if choice == "1":
        data = collect_data_interactive()
    elif choice == "2":
        json_file = input("è¯·è¾“å…¥JSONæ–‡ä»¶è·¯å¾„: ").strip()
        try:
            with open(json_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            print(f"âœ… ä»JSONæ–‡ä»¶åŠ è½½ {len(data)} ä¸ªå­—æ®µ")
        except Exception as e:
            print(f"âŒ åŠ è½½JSONæ–‡ä»¶å¤±è´¥: {{str(e)}}")
            return
    elif choice == "3":
        print("ğŸ‘‹ å†è§ï¼")
        return
    else:
        print("âŒ æ— æ•ˆé€‰æ‹©")
        return

    # å¡«å……æ¨¡æ¿
    output_file = filler.fill_template(data)

    if output_file:
        # ä¿å­˜æ•°æ®
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"ğŸ“„ æ•°æ®å·²ä¿å­˜: {{data_file}}")

if __name__ == "__main__":
    import datetime
    main()
'''

        # ä¿å­˜ç”Ÿæˆçš„ä»£ç 
        output_path = Path(output_file)
        output_path.write_text(code, encoding='utf-8')

        print(f"\nå·²ç”Ÿæˆæ¨¡æ¿å¡«å……ä»£ç : {output_file}")
        print(f"ä½¿ç”¨æ–¹æ³•:")
        print(f"  1. ç¡®ä¿æ¨¡æ¿æ–‡ä»¶ '{template_name}' åœ¨å½“å‰ç›®å½•")
        print(f"  2. è¿è¡Œ: python {output_file}")
        print(f"  3. æŒ‰ç…§æç¤ºå¡«å†™ä¿¡æ¯")

def main():
    """ä¸»å‡½æ•°"""
    print("=" * 60)
    print("  Wordæ¨¡æ¿åˆ†æå·¥å…·")
    print("  ä¸“ä¸ºä¸ªä½“å·¥å•†æˆ·ç™»è®°ç”³è¯·ä¹¦è®¾è®¡")
    print("=" * 60)

    analyzer = TemplateAnalyzer()

    # æŸ¥æ‰¾æ¨¡æ¿æ–‡ä»¶
    template_files = list(Path(".").glob("*.docx"))
    template_files.extend(list(Path(".").glob("*.doc")))

    if not template_files:
        print("æœªæ‰¾åˆ°Wordæ¨¡æ¿æ–‡ä»¶")
        print("æç¤ºï¼šè¯·å°†Wordæ¨¡æ¿æ–‡ä»¶æ”¾åœ¨å½“å‰ç›®å½•")
        return

    print("\næ‰¾åˆ°ä»¥ä¸‹æ¨¡æ¿æ–‡ä»¶ï¼š")
    for i, template_file in enumerate(template_files, 1):
        size_kb = template_file.stat().st_size / 1024
        print(f"{i}. {template_file.name} ({size_kb:.1f} KB)")

    try:
        choice = int(input("\nè¯·é€‰æ‹©è¦åˆ†æçš„æ¨¡æ¿æ–‡ä»¶ç¼–å·: ").strip())
        if 1 <= choice <= len(template_files):
            template_file = template_files[choice - 1]

            # åˆ†ææ¨¡æ¿
            analysis = analyzer.analyze_template(template_file)

            if analysis:
                # ä¿å­˜åˆ†æç»“æœ
                analysis_file = f"{template_file.stem}_analysis.json"
                with open(analysis_file, 'w', encoding='utf-8') as f:
                    json.dump(analysis, f, ensure_ascii=False, indent=2)
                print(f"\nåˆ†æç»“æœå·²ä¿å­˜: {analysis_file}")

                # æ˜¾ç¤ºåˆ†æç»“æœæ‘˜è¦
                print("\nåˆ†æç»“æœæ‘˜è¦ï¼š")
                print(f"  æ¨¡æ¿æ–‡ä»¶: {analysis['template_name']}")
                print(f"  æ–‡ä»¶å¤§å°: {analysis['template_size_kb']:.1f} KB")
                print(f"  æ®µè½æ•°é‡: {analysis['paragraph_count']}")
                print(f"  è¡¨æ ¼æ•°é‡: {analysis['table_count']}")
                print(f"  æ‰¾åˆ°å ä½ç¬¦: {len(analysis['placeholders_found'])} ä¸ª")

                if analysis['placeholders_found']:
                    print("  å…·ä½“å ä½ç¬¦ï¼š")
                    for placeholder in analysis['placeholders_found']:
                        print(f"    - {placeholder}")

                print(f"\nå»ºè®®å­—æ®µ ({len(analysis['suggested_fields'])} ä¸ª)ï¼š")
                for field in analysis['suggested_fields']:
                    required = "å¿…å¡«" if field.get("required", True) else "å¯é€‰"
                    print(f"    - {field['display_name']} ({required}) - {field['description']}")

                # è¯¢é—®æ˜¯å¦ç”Ÿæˆå¡«å……ä»£ç 
                generate_code = input("\næ˜¯å¦ç”Ÿæˆæ¨¡æ¿å¡«å……ä»£ç ï¼Ÿ (y/n): ").strip().lower()
                if generate_code == 'y':
                    code_file = f"{template_file.stem}_filler.py"
                    analyzer.generate_filling_code(analysis, code_file)

                    print(f"\nå®Œæˆï¼")
                    print(f"  1. åˆ†æç»“æœ: {analysis_file}")
                    print(f"  2. å¡«å……ä»£ç : {code_file}")
                    print(f"\nä¸‹ä¸€æ­¥ï¼š")
                    print(f"  è¿è¡Œ: python {code_file}")
                    print(f"  æŒ‰ç…§æç¤ºå¡«å†™ä¿¡æ¯ï¼Œè‡ªåŠ¨ç”Ÿæˆå¡«å……åçš„Wordæ–‡æ¡£")

        else:
            print("æ— æ•ˆçš„é€‰æ‹©")
    except ValueError:
        print("è¯·è¾“å…¥æœ‰æ•ˆçš„æ•°å­—")
    except Exception as e:
        print(f"åˆ†æè¿‡ç¨‹ä¸­å‡ºé”™: {str(e)}")

if __name__ == "__main__":
    import datetime
    main()