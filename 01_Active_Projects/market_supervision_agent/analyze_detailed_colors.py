#!/usr/bin/env python3
"""
详细分析李奕凤模板的字体颜色
"""

from docx import Document

def analyze_colors():
    """详细分析字体颜色"""
    template_file = "（李奕凤）个体工商户开业登记申请书（模板）.docx"
    doc = Document(template_file)

    print("=" * 70)
    print("详细分析字体颜色")
    print("=" * 70)

    for table_idx, table in enumerate(doc.tables):
        print(f"\n=== 表格 {table_idx} ===")

        for row_idx, row in enumerate(table.rows[:5]):  # 只显示前5行
            print(f"\n行 {row_idx}:")

            for cell_idx, cell in enumerate(row.cells):
                print(f"  列 {cell_idx}:")

                for para_idx, paragraph in enumerate(cell.paragraphs):
                    if paragraph.text.strip():
                        for run_idx, run in enumerate(paragraph.runs):
                            text = run.text[:30]

                            # 获取颜色信息
                            color_info = "默认颜色"
                            if run.font.color:
                                if run.font.color.rgb:
                                    r, g, b = run.font.color.rgb
                                    color_info = f"RGB({r},{g},{b})"

                                    # 判断颜色类型
                                    if r > 200 and g < 100 and b < 100:
                                        color_type = "【红色】"
                                    elif r < 50 and g < 50 and b < 50:
                                        color_type = "【黑色】"
                                    else:
                                        color_type = ""
                                    color_info = f"{color_info} {color_type}"
                                elif run.font.color.theme_color:
                                    color_info = f"主题色{run.font.color.theme_color}"

                            print(f"    Run{run_idx}: '{text}' - {color_info}")

if __name__ == "__main__":
    analyze_colors()
