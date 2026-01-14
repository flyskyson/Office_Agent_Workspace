#!/usr/bin/env python3
"""
创建干净的个体工商户开业登记申请书Word模板
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

def create_clean_template():
    """创建干净的Word模板"""

    # 创建新文档
    doc = Document()

    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 添加标题
    title = doc.add_heading('个体工商户开业登记申请书', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加说明文字
    instruction = doc.add_paragraph()
    instruction.alignment = WD_ALIGN_PARAGRAPH.CENTER
    instruction.add_run('注：请将"填写"部分替换为实际信息，然后删除括号内的说明文字')
    instruction.runs[0].font.size = Pt(10)
    instruction.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    # 添加空行
    doc.add_paragraph()

    # 基本信息
    doc.add_heading('一、基本信息', level=2)

    # 创建表格
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'

    # 设置表格内容
    basic_info = [
        ('个体工商户名称', '{{个体工商户名称}}（填写名称）'),
        ('经营者姓名', '{{经营者姓名}}（填写姓名）'),
        ('联系电话', '{{联系电话}}（填写电话）'),
        ('电子邮箱', '{{电子邮箱}}（填写邮箱，可选）'),
        ('经营场所', '{{经营场所}}（填写详细地址）'),
        ('邮政编码', '{{邮政编码}}（填写邮编，可选）'),
        ('从业人数', '{{从业人数}}（填写人数，可选）'),
        ('注册资金', '{{注册资金}}（填写金额，单位：元）')
    ]

    for i, (label, value) in enumerate(basic_info):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

        # 设置单元格样式
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(12)
            cell.paragraphs[0].runs[0].font.name = '宋体'

    doc.add_paragraph()

    # 经营者信息
    doc.add_heading('二、经营者信息', level=2)

    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'

    operator_info = [
        ('身份证号码', '{{身份证号码}}（填写身份证号）'),
        ('性别', '{{性别}}（男/女）'),
        ('民族', '{{民族}}（填写民族，可选）'),
        ('政治面貌', '{{政治面貌}}（填写政治面貌，可选）'),
        ('文化程度', '{{文化程度}}（填写文化程度，可选）')
    ]

    for i, (label, value) in enumerate(operator_info):
        row = table2.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(12)

    doc.add_paragraph()

    # 经营范围
    doc.add_heading('三、经营范围', level=2)

    table3 = doc.add_table(rows=2, cols=1)
    table3.style = 'Table Grid'

    row = table3.rows[0]
    row.cells[0].text = '经营范围'
    row = table3.rows[1]
    row.cells[0].text = '{{经营范围}}（填写详细的经营范围）'

    doc.add_paragraph()

    # 经营期限
    doc.add_heading('四、经营期限', level=2)

    table4 = doc.add_table(rows=1, cols=2)
    table4.style = 'Table Grid'

    row = table4.rows[0]
    row.cells[0].text = '经营期限'
    row.cells[1].text = '{{经营期限}}（如：长期 / 2026年01月11日至2036年01月10日）'

    doc.add_paragraph()

    # 声明部分
    doc.add_heading('五、经营者声明', level=2)

    declaration_text = '''本人承诺所填写内容及提交的材料真实、合法、有效，并对申请材料的真实性负责。

如有虚假，愿承担相应的法律责任。'''

    doc.add_paragraph(declaration_text)

    # 签字和日期
    table5 = doc.add_table(rows=2, cols=2)
    table5.style = 'Table Grid'

    row = table5.rows[0]
    row.cells[0].text = '经营者签字：'
    row.cells[1].text = '{{签字日期}}（填写日期）'

    row = table5.rows[1]
    row.cells[0].text = '（签字）'
    row.cells[1].text = '{{申请日期}}（申请日期）'

    # 添加分页
    doc.add_page_break()

    # 附件清单
    doc.add_heading('六、附件清单', level=2)

    attachments = [
        '1. 经营者身份证复印件（正反面）',
        '2. 经营场所使用证明',
        '3. 《个体工商户开业登记申请书》',
        '4. 委托代理人办理的，需提交委托书及代理人身份证复印件',
        '5. 国家法律法规规定提交的其他文件'
    ]

    for item in attachments:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # 注意事项
    doc.add_heading('七、填写说明', level=2)

    instructions = [
        '1. 本表用黑色钢笔或签字笔填写，字迹应清晰',
        '2. 表中"填写"部分为提示性文字，申请时请替换为实际信息',
        '3. 标有"可选"的项目可以根据实际情况填写或不填',
        '4. 经营范围按《国民经济行业分类》标准填写',
        '5. 经营场所应填写详细地址，具体到门牌号'
    ]

    for instruction in instructions:
        doc.add_paragraph(instruction, style='List Number')

    # 保存模板
    template_file = Path("个体工商户开业登记申请书_新版.docx")
    doc.save(str(template_file))

    print(f"成功创建模板: {template_file}")
    print("\n模板特点:")
    print("  - 干净的格式，无特殊字符")
    print("  - 使用 {{字段名}} 作为占位符")
    print("  - 包含所有必要字段")
    print("  - 附带填写说明")

    return template_file

if __name__ == "__main__":
    create_clean_template()