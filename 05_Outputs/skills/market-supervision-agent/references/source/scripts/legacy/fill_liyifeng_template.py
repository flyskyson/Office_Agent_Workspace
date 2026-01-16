#!/usr/bin/env python3
"""
填充李奕凤版模板
- 红色字体 = 字段值（需要填充的用户数据）
- 黑色字体 = 字段名（保留不变）
"""

import json
import datetime
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

def get_font_color(run):
    """获取字体颜色"""
    if run.font.color and run.font.color.rgb:
        return run.font.color.rgb
    return None

def is_red_font(run):
    """判断是否为红色字体"""
    color = get_font_color(run)
    if color:
        # 红色的RGB值接近 (255, 0, 0)
        r, g, b = color
        return r > 200 and g < 100 and b < 100
    return False

def fill_liyifeng_template(data, template_file="（李奕凤）个体工商户开业登记申请书（模板）.docx"):
    """填充李奕凤版模板"""
    try:
        print(f"正在读取模板: {template_file}")
        doc = Document(template_file)

        print("正在智能填充模板（识别红色字体字段）...")

        # 定义字段映射（根据字段名匹配数据）
        field_keywords = {
            "个体工商户名称": "business_name",
            "经营者姓名": "operator_name",
            "联系电话": "phone",
            "电子邮箱": "email",
            "经营场所": "business_address",
            "邮政编码": "postal_code",
            "从业人数": "employee_count",
            "注册资金": "registered_capital",
            "身份证号码": "id_card",
            "性别": "gender",
            "民族": "nation",
            "政治面貌": "political_status",
            "文化程度": "education",
            "经营范围": "business_scope",
            "经营期限": "operation_period"
        }

        filled_count = 0

        # 遍历所有表格
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # 检查单元格中的所有段落和run
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            # 检查是否为红色字体（需要填充的字段值）
                            if is_red_font(run):
                                run_text = run.text.strip()

                                # 查找对应的字段
                                for keyword, data_key in field_keywords.items():
                                    if keyword in run_text or run_text in ["{{", "}}", "模板", "示例"]:
                                        # 获取对应的值
                                        if data_key in data:
                                            value = str(data.get(data_key, ""))
                                            # 清理非法字符
                                            clean_value = ''.join(c for c in value if ord(c) < 0xD800 or ord(c) > 0xDFFF)

                                            # 只替换红色字体的内容
                                            run.text = clean_value
                                            filled_count += 1
                                            print(f"  已填充红色字段 [{keyword}]: {clean_value}")
                                            break

        print(f"\n完成 {filled_count} 处红色字段填充")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_开业登记_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")
        print(f"[说明] 李奕凤版模板，红色字体已替换为实际数据")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] {data_file}")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数"""
    # 测试数据
    test_data = {
        'business_name': '吴九药店',
        'operator_name': '吴九',
        'phone': '13444444444',
        'email': 'wujiu@test.com',
        'business_address': '武汉市江汉区中山大道888号',
        'postal_code': '430000',
        'employee_count': '12',
        'registered_capital': '500000',
        'id_card': '420101198606065555',
        'gender': '女',
        'nation': '汉族',
        'political_status': '群众',
        'education': '本科',
        'business_scope': '药品零售；医疗器械销售',
        'operation_period': '长期'
    }

    print("=" * 70)
    print("李奕凤版模板填充工具")
    print("红色字体 = 字段值（自动填充）")
    print("黑色字体 = 字段名（保留不变）")
    print("=" * 70)

    print("\n测试数据：")
    for key, value in test_data.items():
        print(f"  {key}: {value}")

    result = fill_liyifeng_template(test_data)

    if result:
        print("\n" + "=" * 70)
        print("测试成功！")
        print(f"生成的文件: {result}")
        print("=" * 70)

        # 用WPS打开
        print("\n正在用WPS打开文档...")
        import subprocess
        subprocess.Popen(['cmd', '/c', 'start', '', str(result)], shell=True)

if __name__ == "__main__":
    main()
