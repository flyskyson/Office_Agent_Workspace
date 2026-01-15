#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修复模板字体样式
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def fix_template_font():
    """修复模板字体"""

    # 读取模板
    doc = Document('（李奕凤）个体工商户开业登记申请书（Jinja2模板）.docx')

    print("正在调整模板中的字体样式...")

    # 遍历所有段落和表格，设置字体为宋体
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if run.font.size is None:
                run.font.size = Pt(10.5)

    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        if run.font.size is None:
                            run.font.size = Pt(10.5)

    # 保存为新文件
    output_file = '（李奕凤）个体工商户开业登记申请书（Jinja2模板-字体修复）.docx'
    doc.save(output_file)

    print(f"[OK] 字体样式已调整完成！")
    print(f"[文件] {output_file}")
    print("所有文字已设置为宋体 10.5号")

if __name__ == "__main__":
    fix_template_font()
