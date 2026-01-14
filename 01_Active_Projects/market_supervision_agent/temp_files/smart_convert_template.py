#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能模板转换工具 v2 - 基于实际红色字体内容进行转换
"""

from docx import Document
import json

def is_red_font(run):
    """判断是否为红色字体"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        return r == 255 and g == 0 and b == 0
    return False

def smart_convert_template():
    """智能转换模板"""

    print("=" * 70)
    print("智能模板转换工具 v2")
    print("=" * 70)

    # 读取原模板
    src = Document(r'c:\Users\flyskyson\Desktop\（李奕凤）个体工商户开业登记申请书（模板-待修改）.docx')

    # 基于实际红色字体内容的映射
    red_font_mappings = {
        # 个体工商户名称相关
        "兴业县蒲塘镇源兴快餐店（个体工商户）（个体工商户）": "{{business_name}}",
        "兴业县蒲塘镇源兴快餐店": "{{business_name}}",

        # 经营者信息
        "李奕凤": "{{operator_name}}",
        "女": "{{gender}}",
        "汉": "{{nation}}",
        "群众": "{{political_status}}",

        # 联系方式
        "18207759518": "{{phone}}",

        # 地址信息
        "广西壮族自治区玉林市兴业县蒲塘镇蒲安路238号（蒲塘镇客运站旁）": "{{business_address}}",
        "平塘一队5号": "{{business_address}}",  # 住所地址

        # 身份证
        "452501196705095123": "{{id_card}}",

        # 数字
        "2": "{{employee_count}}",
        "1": "",  # 这个可能是复选框，暂时忽略
        "10000": "{{registered_capital}}",
        "537820": "{{postal_code}}",  # 邮编（但这个可能是绿色常量）

        # 经营范围相关
        "许可项目：": "",  # 标题，保持不变
        "小餐饮": "{{business_scope_licensed}}",
        "（依法须经批准的项目，经相关部门批准后方可开展经营活动，具体经营项目以相关部门批准文件或许可证件为准": "",
        "一般项目：": "",
        "食品销售（仅销售预包装食品）": "{{business_scope_general}}",
        "（除依法须经批准的项目外，凭营业执照依法自主开展经营活动）": ""
    }

    converted_count = 0
    skipped_count = 0

    print("\n正在转换...")

    # 处理表格中的红色字体
    for table_idx, table in enumerate(src.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for run in para.runs:
                        if is_red_font(run):
                            text = run.text.strip()

                            # 查找匹配的映射
                            if text in red_font_mappings:
                                new_text = red_font_mappings[text]

                                if new_text:  # 如果有替换内容
                                    run.text = new_text
                                    converted_count += 1
                                    print(f"  [OK] Table{table_idx} Row{row_idx}: '{text[:20]}' -> '{new_text}'")
                                else:  # 删除空内容
                                    run.text = ""
                                    skipped_count += 1
                                    print(f"  [SKIP] Table{table_idx} Row{row_idx}: '{text[:20]}' (删除)")

    # 保存转换后的模板
    output_file = r'c:\Users\flyskyson\Desktop\（李奕凤）个体工商户开业登记申请书（Jinja2模板）.docx'
    src.save(output_file)

    print("\n" + "=" * 70)
    print("转换完成！")
    print("=" * 70)
    print(f"\n输出文件: {output_file}")
    print(f"成功转换: {converted_count} 处")
    print(f"跳过/删除: {skipped_count} 处")

    print("\n下一步:")
    print("1. 在WPS中打开转换后的模板检查")
    print("2. 如有需要，手动调整未转换的部分")
    print("3. 使用以下命令测试:")
    print("   python jinja2_filler.py --template \"Jinja2模板.docx\" --test")
    print("=" * 70)

if __name__ == "__main__":
    smart_convert_template()
