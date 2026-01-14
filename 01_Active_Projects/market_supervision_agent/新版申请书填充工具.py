#!/usr/bin/env python3
"""
个体工商户开业登记申请书填充工具 - 新版
配合新版模板使用
"""

import sys
from pathlib import Path
import json
import datetime

try:
    from docx import Document
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)

def check_template():
    """检查模板文件"""
    template_file = "个体工商户开业登记申请书_新版.docx"
    if not Path(template_file).exists():
        print(f"错误: 模板文件不存在")
        print(f"请确保 '{template_file}' 在当前目录")
        print("\n提示: 运行 'python create_clean_template.py' 创建新模板")
        return False
    return True

def collect_data():
    """收集用户数据"""
    print("\n请填写个体工商户开业登记信息")
    print("=" * 60)
    print("说明: 带*的为必填项，其他可以按Enter跳过")

    data = {}

    # 基本信息
    print("\n[基本信息]")
    data["business_name"] = input("*个体工商户名称: ").strip()
    data["operator_name"] = input("*经营者姓名: ").strip()
    data["phone"] = input("*联系电话: ").strip()
    data["email"] = input("  电子邮箱（可选）: ").strip() or "未填写"
    data["business_address"] = input("*经营场所: ").strip()
    data["postal_code"] = input("  邮政编码（可选）: ").strip() or "未填写"
    data["employee_count"] = input("  从业人数（可选）: ").strip() or "未填写"
    data["registered_capital"] = input("*注册资金（元）: ").strip()

    # 经营者信息
    print("\n[经营者信息]")
    data["id_card"] = input("*身份证号码: ").strip()
    data["gender"] = input("*性别（男/女）: ").strip()
    data["nation"] = input("  民族（可选）: ").strip() or "未填写"
    data["political_status"] = input("  政治面貌（可选）: ").strip() or "未填写"
    data["education"] = input("  文化程度（可选）: ").strip() or "未填写"

    # 经营信息
    print("\n[经营信息]")
    data["business_scope"] = input("*经营范围: ").strip()
    data["operation_period"] = input("*经营期限: ").strip() or "长期"

    # 自动添加日期
    data["application_date"] = datetime.datetime.now().strftime("%Y年%m月%d日")
    data["sign_date"] = data["application_date"]

    return data

def fill_template(data):
    """填充模板"""
    try:
        # 读取模板
        template_file = "个体工商户开业登记申请书_新版.docx"
        doc = Document(template_file)

        print("\n正在填充模板...")

        # 创建替换映射 - 使用双花括号格式
        replacements = {
            # 基本信息
            "{{个体工商户名称}}": data.get("business_name", ""),
            "{{经营者姓名}}": data.get("operator_name", ""),
            "{{联系电话}}": data.get("phone", ""),
            "{{电子邮箱}}": data.get("email", ""),
            "{{经营场所}}": data.get("business_address", ""),
            "{{邮政编码}}": data.get("postal_code", ""),
            "{{从业人数}}": data.get("employee_count", ""),
            "{{注册资金}}": data.get("registered_capital", ""),

            # 经营者信息
            "{{身份证号码}}": data.get("id_card", ""),
            "{{性别}}": data.get("gender", ""),
            "{{民族}}": data.get("nation", ""),
            "{{政治面貌}}": data.get("political_status", ""),
            "{{文化程度}}": data.get("education", ""),

            # 经营信息
            "{{经营范围}}": data.get("business_scope", ""),
            "{{经营期限}}": data.get("operation_period", ""),

            # 其他
            "{{签字日期}}": data.get("sign_date", ""),
            "{{申请日期}}": data.get("application_date", "")
        }

        # 在段落中替换
        replaced_count = 0
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
                    replaced_count += 1

        # 在表格中替换
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old_text, new_text in replacements.items():
                        if old_text in cell.text:
                            cell.text = cell.text.replace(old_text, new_text)
                            replaced_count += 1

        print(f"完成 {replaced_count} 处替换")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_开业登记申请书_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] 位置: {output_file}")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] 文件: {data_file}")

        print(f"\n[提示] 请打开生成的Word文件检查")
        print(f"       如需修改，可以直接在Word中编辑")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def load_from_json():
    """从JSON文件加载数据"""
    json_file = input("请输入JSON文件路径: ").strip()
    if not Path(json_file).exists():
        print(f"文件不存在: {json_file}")
        return None

    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        print(f"从 {json_file} 加载数据成功")
        return data
    except Exception as e:
        print(f"加载失败: {str(e)}")
        return None

def create_example_data():
    """创建示例数据文件"""
    example = {
        "business_name": "张三小吃店",
        "operator_name": "张三",
        "phone": "13800138000",
        "email": "zhangsan@example.com",
        "business_address": "北京市东城区王府井大街1号",
        "postal_code": "100006",
        "employee_count": "3",
        "registered_capital": "50000",
        "id_card": "110101199001011234",
        "gender": "男",
        "nation": "汉族",
        "political_status": "群众",
        "education": "高中",
        "business_scope": "餐饮服务；小吃店经营",
        "operation_period": "长期",
        "application_date": "2026年01月11日",
        "sign_date": "2026年01月11日"
    }

    example_file = Path("example_data.json")
    with open(example_file, 'w', encoding='utf-8') as f:
        json.dump(example, f, ensure_ascii=False, indent=2)

    print(f"示例数据文件已创建: {example_file}")
    print("可以使用这个文件测试批量填充功能")

def main():
    """主函数"""
    print("=" * 60)
    print("个体工商户开业登记申请书填充工具 - 新版")
    print("=" * 60)

    # 检查模板
    if not check_template():
        return

    print("\n请选择操作:")
    print("1. 交互式填写")
    print("2. 从JSON文件加载")
    print("3. 创建示例数据文件")
    print("4. 退出")

    choice = input("\n请输入选项 (1-4): ").strip()

    if choice == "1":
        data = collect_data()
    elif choice == "2":
        data = load_from_json()
        if not data:
            return
    elif choice == "3":
        create_example_data()
        return
    elif choice == "4":
        print("\n再见！")
        return
    else:
        print("\n无效选项")
        return

    # 填充模板
    result = fill_template(data)

    if result:
        print("\n" + "=" * 60)
        print("完成！")
        print(f"生成的文件在 output/ 目录")
        print("=" * 60)

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n程序已取消")
    except Exception as e:
        print(f"\n程序错误: {str(e)}")