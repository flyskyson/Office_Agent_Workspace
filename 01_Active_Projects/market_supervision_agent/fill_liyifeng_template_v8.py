#!/usr/bin/env python3
"""
智能填充李奕凤版模板 v8 - 跨单元格关联版
- 删除所有红色字体内容
- 识别相邻单元格的字段名（左侧单元格）
- 填充实际数据到红色字体位置
"""

import json
import datetime
from pathlib import Path
from docx import Document

def is_red_font(run):
    """判断是否为红色字体 RGB(255,0,0)"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        return r == 255 and g == 0 and b == 0
    return False

def cell_has_red_font(cell):
    """检查单元格是否有红色字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if is_red_font(run):
                return True
    return False

def get_cell_text_clean(cell):
    """获取单元格的纯文本（用于字段名识别）"""
    text_parts = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if not is_red_font(run) and run.text.strip():
                text_parts.append(run.text.strip())
    return "".join(text_parts)

def fill_liyifeng_template_v8(data, template_file="（李奕凤）个体工商户开业登记申请书（模板）.docx"):
    """智能填充李奕凤版模板 v8 - 跨单元格关联"""
    try:
        print(f"正在读取模板: {template_file}")
        doc = Document(template_file)

        print("正在处理：识别相邻单元格，删除红色字体，填充数据...")

        # 定义字段识别规则
        field_patterns = [
            (["个体工商户名称"], "business_name"),
            (["经营者姓名", "经营者"], "operator_name"),
            (["联系电话", "电话"], "phone"),
            (["电子邮箱", "邮箱"], "email"),
            (["经营场所", "住所", "地址"], "business_address"),
            (["邮政编码", "邮编"], "postal_code"),
            (["从业人数", "人数"], "employee_count"),
            (["注册资金", "资金"], "registered_capital"),
            (["身份证号码", "身份证"], "id_card"),
            (["性别"], "gender"),
            (["民族"], "nation"),
            (["政治面貌", "政治"], "political_status"),
            (["文化程度", "学历"], "education"),
            (["经营范围"], "business_scope"),
            (["经营期限", "期限"], "operation_period")
        ]

        deleted_count = 0
        filled_count = 0
        filled_fields = set()

        # 遍历所有表格
        for table_idx, table in enumerate(doc.tables):
            print(f"\n处理表格 {table_idx}...")

            for row_idx, row in enumerate(table.rows):
                # 先找出这一行有红色字体的单元格
                red_cells = []
                for cell_idx, cell in enumerate(row.cells):
                    if cell_has_red_font(cell):
                        red_cells.append(cell_idx)

                # 对每个有红色字体的单元格，查找左侧单元格的字段名
                for cell_idx in red_cells:
                    cell = row.cells[cell_idx]

                    # 查找左侧单元格的字段名
                    field_name_text = ""
                    if cell_idx > 0:
                        left_cell = row.cells[cell_idx - 1]
                        field_name_text = get_cell_text_clean(left_cell)

                    # 如果左侧没有，尝试当前单元格删除红色后的文本
                    if not field_name_text:
                        field_name_text = get_cell_text_clean(cell)

                    # 查找匹配的字段
                    matched_key = None
                    for field_names, data_key in field_patterns:
                        if data_key not in filled_fields:  # 未填充过的字段
                            for field_name in field_names:
                                if field_name in field_name_text:
                                    matched_key = data_key
                                    filled_fields.add(data_key)
                                    break
                            if matched_key:
                                break

                    # 删除红色字体
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if is_red_font(run):
                                run.text = ""
                                deleted_count += 1

                    # 如果找到匹配，添加数据
                    if matched_key and matched_key in data:
                        value = str(data[matched_key])
                        if value and cell.paragraphs:
                            last_para = cell.paragraphs[-1]

                            # 获取格式
                            font_name = None
                            font_size = None
                            if last_para.runs:
                                last_run = last_para.runs[-1]
                                if last_run.font.name:
                                    font_name = last_run.font.name
                                if last_run.font.size:
                                    font_size = last_run.font.size

                            # 添加新run
                            new_run = last_para.add_run(value)
                            if font_name:
                                new_run.font.name = font_name
                            if font_size:
                                new_run.font.size = font_size

                            filled_count += 1
                            print(f"  [OK] {matched_key}: {value}")

        print(f"\n完成统计:")
        print(f"  删除红色字体: {deleted_count} 处")
        print(f"  填充数据: {filled_count} 处")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_开业登记_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")
        print(f"[说明] 已删除红色字体示例数据，填充实际信息")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] {data_file}")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数"""
    # 测试数据
    test_data = {
        'business_name': '张三便利店',
        'operator_name': '张三',
        'phone': '13111111111',
        'email': 'zhangsan@test.com',
        'business_address': '北京市朝阳区建国路88号',
        'postal_code': '100000',
        'employee_count': '3',
        'registered_capital': '50000',
        'id_card': '110101198001011111',
        'gender': '男',
        'nation': '汉族',
        'political_status': '群众',
        'education': '高中',
        'business_scope': '日用百货；食品销售',
        'operation_period': '长期'
    }

    print("=" * 70)
    print("李奕凤版模板填充工具 v8 - 跨单元格关联版")
    print("识别相邻单元格字段名，删除红色字体，填充数据")
    print("=" * 70)

    print("\n测试数据：")
    for key, value in test_data.items():
        print(f"  {key}: {value}")

    result = fill_liyifeng_template_v8(test_data)

    if result:
        print("\n" + "=" * 70)
        print("测试成功！")
        print(f"生成的文件: {result}")
        print("=" * 70)

        # 用WPS打开
        print("\n正在用WPS打开文档...")
        import subprocess
        subprocess.Popen(['cmd', '/c', 'start', '', str(result)], shell=True)

if __name__ == "__main__":
    main()
