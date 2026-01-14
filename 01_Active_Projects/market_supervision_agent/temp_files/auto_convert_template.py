#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能模板转换工具 - 自动将红色字体转换为 Jinja2 变量
"""

from docx import Document
import json

def is_red_font(run):
    """判断是否为红色字体"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        return r == 255 and g == 0 and b == 0
    return False

def auto_convert_template():
    """自动转换模板"""

    print("=" * 70)
    print("智能模板转换工具")
    print("=" * 70)
    print("\n正在读取原模板...")

    # 读取原模板
    src = Document(r'c:\Users\flyskyson\Desktop\（李奕凤）个体工商户开业登记申请书（模板-待修改）.docx')

    # 变量映射表
    field_mappings = {
        "个体工商户名称": "business_name",
        "经营者姓名": "operator_name",
        "经营者": "operator_name",
        "联系电话": "phone",
        "电话": "phone",
        "电子邮箱": "email",
        "邮箱": "email",
        "经营场所": "business_address",
        "住所": "business_address",
        "地址": "business_address",
        "邮政编码": "postal_code",
        "邮编": "postal_code",
        "从业人数": "employee_count",
        "人数": "employee_count",
        "注册资金": "registered_capital",
        "资金": "registered_capital",
        "身份证号码": "id_card",
        "身份证": "id_card",
        "性别": "gender",
        "民族": "nation",
        "政治面貌": "political_status",
        "政治": "political_status",
        "文化程度": "education",
        "学历": "education",
        "经营范围": "business_scope",
        "经营期限": "operation_period",
        "期限": "operation_period"
    }

    converted_count = 0
    green_preserved = 0

    print("正在转换红色字体为 Jinja2 变量...")

    # 处理段落
    for para in src.paragraphs:
        for run in para.runs:
            if is_red_font(run):
                # 检查是否可以映射
                text = run.text.strip()
                for cn_name, var_name in field_mappings.items():
                    if cn_name in text or text in cn_name:
                        run.text = "{{" + var_name + "}}"
                        converted_count += 1
                        print(f"  [OK] Convert: {text[:20]} -> {{{{var_name}}}}")
                        break

    # 处理表格
    for table_idx, table in enumerate(src.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for run in para.runs:
                        if is_red_font(run):
                            text = run.text.strip()
                            for cn_name, var_name in field_mappings.items():
                                if cn_name in text or text in cn_name:
                                    run.text = "{{" + var_name + "}}"
                                    converted_count += 1
                                    print(f"  [OK] Table{table_idx} Row{row_idx}: {text[:20]} -> {{{{var_name}}}}")
                                    break

    # 保存转换后的模板
    output_file = r'c:\Users\flyskyson\Desktop\（李奕凤）个体工商户开业登记申请书（Jinja2模板）.docx'
    src.save(output_file)

    print("\n" + "=" * 70)
    print("转换完成！")
    print("=" * 70)
    print(f"\n输出文件: {output_file}")
    print(f"转换数量: {converted_count} 处")
    print(f"保留绿色: {green_preserved} 处")
    print("\n下一步:")
    print("1. 在WPS中打开转换后的模板")
    print("2. 检查转换结果，手动调整未转换的部分")
    print("3. 使用以下命令测试:")
    print("   python jinja2_filler.py --template \"Jinja2模板.docx\" --test")
    print("=" * 70)

if __name__ == "__main__":
    auto_convert_template()
