#!/usr/bin/env python3
"""
个体工商户Word申请表生成系统
专为市场监管工作人员设计 - 支持Word模板
"""

import os
import json
import datetime
from pathlib import Path
from typing import Dict, Any, List
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ 需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)

class WordApplicationGenerator:
    """Word申请表生成器 - 支持模板填充"""

    def __init__(self, template_dir: str = "templates"):
        """
        初始化Word申请表生成器

        Args:
            template_dir: 模板文件目录
        """
        self.template_dir = Path(template_dir)
        self.template_dir.mkdir(exist_ok=True)

        print("=" * 60)
        print("  个体工商户Word申请表生成系统")
        print("  版本 1.0 - 支持Word模板")
        print("=" * 60)

    def show_menu(self):
        """显示主菜单"""
        print("\n📋 请选择要生成的申请表类型：")
        print("1. 个体工商户设立登记申请书 (Word)")
        print("2. 个体工商户变更登记申请书 (Word)")
        print("3. 个体工商户注销登记申请书 (Word)")
        print("4. 个体工商户年度报告表 (Word)")
        print("5. 使用现有Word模板生成")
        print("6. 查看可用模板")
        print("7. 退出系统")

        choice = input("\n请输入选项编号 (1-7): ").strip()
        return choice

    def generate_registration_word(self):
        """生成设立登记申请书Word文档"""
        print("\n📝 生成个体工商户设立登记申请书 (Word)")
        print("请填写以下信息（按Enter跳过可选字段）：")

        data = {
            "business_name": input("个体工商户名称: "),
            "operator_name": input("经营者姓名: "),
            "gender": input("性别（男/女）: "),
            "id_card": input("身份证号码: "),
            "phone": input("联系电话: "),
            "email": input("电子邮箱（可选）: ") or "未提供",
            "business_address": input("经营场所: "),
            "postal_code": input("邮政编码（可选）: ") or "未提供",
            "business_area": input("经营面积（平方米，可选）: ") or "未提供",
            "business_scope": input("经营范围: "),
            "main_business": input("主营业务（可选）: ") or "未提供",
            "business_type": input("行业类型（可选）: ") or "未提供",
            "registered_capital": input("资金数额（元）: "),
            "capital_source": input("资金来源（可选）: ") or "自有资金",
            "investment_form": input("出资形式（可选）: ") or "货币",
            "operation_period": input("经营期限（可选）: ") or "长期",
            "application_date": datetime.datetime.now().strftime("%Y年%m月%d日")
        }

        # 创建Word文档
        doc = Document()

        # 添加标题
        title = doc.add_heading('个体工商户设立登记申请书', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加申请事项
        doc.add_heading('申请事项：个体工商户设立登记', level=1)

        # 添加基本信息
        doc.add_heading('一、基本信息', level=2)
        self._add_bullet_point(doc, f"1. 个体工商户名称：{data['business_name']}")
        self._add_bullet_point(doc, f"2. 经营者姓名：{data['operator_name']}")
        self._add_bullet_point(doc, f"3. 性别：{data['gender']}")
        self._add_bullet_point(doc, f"4. 身份证号码：{data['id_card']}")
        self._add_bullet_point(doc, f"5. 联系电话：{data['phone']}")
        self._add_bullet_point(doc, f"6. 电子邮箱：{data['email']}")

        # 添加经营信息
        doc.add_heading('二、经营信息', level=2)
        self._add_bullet_point(doc, f"1. 经营场所：{data['business_address']}")
        self._add_bullet_point(doc, f"2. 邮政编码：{data['postal_code']}")
        self._add_bullet_point(doc, f"3. 经营面积：{data['business_area']}平方米")
        self._add_bullet_point(doc, f"4. 经营范围：{data['business_scope']}")
        self._add_bullet_point(doc, f"5. 主营业务：{data['main_business']}")
        self._add_bullet_point(doc, f"6. 行业类型：{data['business_type']}")

        # 添加资金信息
        doc.add_heading('三、资金信息', level=2)
        self._add_bullet_point(doc, f"1. 资金数额：{data['registered_capital']}元")
        self._add_bullet_point(doc, f"2. 资金来源：{data['capital_source']}")
        self._add_bullet_point(doc, f"3. 出资形式：{data['investment_form']}")
        self._add_bullet_point(doc, f"4. 经营期限：{data['operation_period']}")

        # 添加经营者声明
        doc.add_heading('四、经营者声明', level=2)
        declaration = doc.add_paragraph()
        declaration.add_run("本人承诺所填写内容及提交的材料真实、合法、有效，并对申请材料的真实性负责。")

        # 添加签字和日期
        doc.add_paragraph()
        sign_line = doc.add_paragraph()
        sign_line.add_run("经营者签字：___________________")
        sign_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        date_line = doc.add_paragraph()
        date_line.add_run(f"申请日期：{data['application_date']}")
        date_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 添加附件清单
        doc.add_heading('五、附件清单', level=2)
        self._add_bullet_point(doc, "1. 经营者身份证复印件")
        self._add_bullet_point(doc, "2. 经营场所使用证明")
        self._add_bullet_point(doc, "3. 其他相关材料")

        # 保存文件
        output_dir = Path("generated_applications") / "word"
        output_dir.mkdir(parents=True, exist_ok=True)

        filename = f"{data['business_name']}_设立登记申请书_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_file = output_dir / filename

        doc.save(str(output_file))

        # 保存数据
        data_file = output_dir / f"{filename}.json"
        data_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')

        print(f"\n✅ Word申请表生成成功！")
        print(f"📁 文件保存位置: {output_file}")
        print(f"📄 文件名: {filename}")

        return output_file

    def _add_bullet_point(self, doc, text):
        """添加带项目符号的段落"""
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(text)

    def use_existing_template(self):
        """使用现有Word模板生成"""
        print("\n📂 使用现有Word模板生成申请表")

        # 查找现有的Word模板
        template_files = list(Path(".").glob("*.docx"))
        template_files.extend(list(Path(".").glob("*.doc")))

        if not template_files:
            print("❌ 未找到Word模板文件")
            print("请将Word模板文件放在当前目录")
            return

        print("\n📋 找到以下模板文件：")
        for i, template_file in enumerate(template_files, 1):
            print(f"{i}. {template_file.name}")

        try:
            choice = int(input("\n请选择模板文件编号: ").strip())
            if 1 <= choice <= len(template_files):
                template_file = template_files[choice - 1]
                print(f"✅ 选择模板: {template_file.name}")

                # 询问用户要填写的信息
                print("\n📝 请填写模板需要的信息：")
                data = {}

                # 常见字段
                common_fields = [
                    ("个体工商户名称", "business_name"),
                    ("经营者姓名", "operator_name"),
                    ("身份证号码", "id_card"),
                    ("联系电话", "phone"),
                    ("经营场所", "business_address"),
                    ("经营范围", "business_scope"),
                    ("资金数额", "registered_capital"),
                    ("申请日期", "application_date")
                ]

                for display_name, field_name in common_fields:
                    value = input(f"{display_name}: ").strip()
                    if value:
                        data[field_name] = value

                # 尝试读取模板并替换内容
                self._fill_word_template(template_file, data)

            else:
                print("❌ 无效的选择")
        except ValueError:
            print("❌ 请输入有效的数字")

    def _fill_word_template(self, template_path: Path, data: Dict[str, str]):
        """填充Word模板"""
        try:
            # 读取模板
            doc = Document(str(template_path))

            # 简单的文本替换（实际应用中需要更复杂的逻辑）
            # 这里我们假设模板中有特定的占位符，如 {{business_name}}
            replacements = {
                "{{business_name}}": data.get("business_name", ""),
                "{{operator_name}}": data.get("operator_name", ""),
                "{{id_card}}": data.get("id_card", ""),
                "{{phone}}": data.get("phone", ""),
                "{{business_address}}": data.get("business_address", ""),
                "{{business_scope}}": data.get("business_scope", ""),
                "{{registered_capital}}": data.get("registered_capital", ""),
                "{{application_date}}": data.get("application_date", datetime.datetime.now().strftime("%Y年%m月%d日"))
            }

            # 遍历所有段落进行替换
            for paragraph in doc.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, new_text)

            # 遍历所有表格进行替换
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for old_text, new_text in replacements.items():
                            if old_text in cell.text:
                                cell.text = cell.text.replace(old_text, new_text)

            # 保存新文件
            output_dir = Path("generated_applications") / "word"
            output_dir.mkdir(parents=True, exist_ok=True)

            filename = f"filled_{template_path.stem}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_file = output_dir / filename

            doc.save(str(output_file))

            print(f"\n✅ 模板填充成功！")
            print(f"📁 文件保存位置: {output_file}")
            print(f"📄 文件名: {filename}")

            # 保存数据
            data_file = output_dir / f"{filename}.json"
            data_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')

        except Exception as e:
            print(f"❌ 填充模板失败: {str(e)}")
            print("💡 提示：")
            print("1. 确保模板文件没有损坏")
            print("2. 确保有足够的权限")
            print("3. 模板可能需要手动调整格式")

    def view_templates(self):
        """查看可用模板"""
        print("\n📚 可用模板文件：")

        # 查找所有模板文件
        template_files = []

        # Word模板
        word_templates = list(Path(".").glob("*.docx"))
        word_templates.extend(list(Path(".").glob("*.doc")))

        # 文本模板
        text_templates = list(self.template_dir.glob("*.txt"))

        if word_templates:
            print("\n📄 Word模板：")
            for i, template in enumerate(word_templates, 1):
                size_kb = template.stat().st_size / 1024
                print(f"  {i}. {template.name} ({size_kb:.1f} KB)")
                template_files.append(template)

        if text_templates:
            print("\n📝 文本模板：")
            for i, template in enumerate(text_templates, len(word_templates) + 1):
                size_kb = template.stat().st_size / 1024
                print(f"  {i}. {template.name} ({size_kb:.1f} KB)")
                template_files.append(template)

        if not template_files:
            print("❌ 未找到模板文件")
            print("💡 提示：将模板文件放在当前目录或 templates/ 目录")

        return template_files

    def run(self):
        """运行主程序"""
        print("\n🎉 欢迎使用个体工商户Word申请表生成系统！")
        print("本系统支持从模板生成规范的Word格式申请表。")

        while True:
            choice = self.show_menu()

            if choice == "1":
                self.generate_registration_word()
            elif choice == "2":
                print("变更登记Word功能开发中...")
                # self.generate_change_word()
            elif choice == "3":
                print("注销登记Word功能开发中...")
                # self.generate_cancellation_word()
            elif choice == "4":
                print("年报Word功能开发中...")
                # self.generate_annual_report_word()
            elif choice == "5":
                self.use_existing_template()
            elif choice == "6":
                self.view_templates()
            elif choice == "7":
                print("\n👋 感谢使用，再见！")
                break
            else:
                print("❌ 无效的选择，请重新输入")

            # 询问是否继续
            if choice != "7":
                continue_choice = input("\n是否继续生成其他申请表？ (y/n): ").strip().lower()
                if continue_choice != 'y':
                    print("\n👋 感谢使用，再见！")
                    break

def main():
    """主函数"""
    try:
        generator = WordApplicationGenerator()
        generator.run()
    except KeyboardInterrupt:
        print("\n\n👋 程序已取消")
    except Exception as e:
        print(f"\n❌ 程序运行出错: {str(e)}")
        import traceback
        traceback.print_exc()
        print("\n💡 请检查：")
        print("1. 确保已安装 python-docx: pip install python-docx")
        print("2. 确保Word模板文件格式正确")
        print("3. 确保有文件写入权限")

if __name__ == "__main__":
    main()