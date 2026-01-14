#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析模板中所有的颜色
"""

from docx import Document
import json
from collections import defaultdict

def analyze_all_colors():
    """分析所有颜色"""
    doc = Document(r'c:\Users\flyskyson\Desktop\（李奕凤）个体工商户开业登记申请书（模板-待修改）.docx')

    color_stats = defaultdict(list)
    all_colors = set()

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text.strip()
                        if not text:
                            continue

                        # 获取颜色
                        if run.font.color and run.font.color.rgb:
                            r, g, b = run.font.color.rgb
                            color_key = f'RGB({r},{g},{b})'
                            all_colors.add(color_key)

                            # 判断颜色类型
                            if r == 255 and g == 0 and b == 0:
                                color_type = '红色'
                            elif r == 0 and g == 255 and b == 0:
                                color_type = '亮绿色'
                            elif r == 0 and g == 128 and b == 0:
                                color_type = '深绿色'
                            elif g > r and g > b:
                                color_type = '绿色系'
                            else:
                                color_type = '其他'

                            color_stats[color_key].append({
                                'text': text[:50],  # 限制长度
                                'type': color_type,
                                'location': f'表{table_idx}-行{row_idx}-列{cell_idx}'
                            })

    # 打印所有发现的颜色
    print("=" * 70)
    print("模板中所有颜色统计")
    print("=" * 70)

    print(f"\n发现 {len(all_colors)} 种不同的颜色:\n")

    for color in sorted(all_colors):
        items = color_stats[color]
        print(f"\n【{color}】 - {len(items)} 个")
        if items:
            color_type = items[0]['type']
            print(f"  类型: {color_type}")
            print(f"  示例内容:")
            for item in items[:5]:  # 只显示前5个
                print(f"    - {item['text']} ({item['location']})")
            if len(items) > 5:
                print(f"    ... 还有 {len(items) - 5} 个")

    # 保存详细结果
    result = {
        'total_colors': len(all_colors),
        'colors': dict(color_stats)
    }

    output_file = r'c:\Users\flyskyson\Desktop\all_colors_analysis.json'
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"\n详细分析已保存到: {output_file}")

    return result

if __name__ == "__main__":
    analyze_all_colors()
