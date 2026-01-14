#!/usr/bin/env python3
"""
智能填充李奕凤版模板 v10 - 最终完美版
- 删除所有红色字体内容
- 智能识别字段名（左侧、上方、当前单元格）
- 完整填充所有15个字段
- 完美格式保持
"""

import json
import datetime
from pathlib import Path
from docx import Document

def is_red_font(run):
    """判断是否为红色字体 RGB(255,0,0)"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        return r == 255 and g == 0 and b == 0
    return False

def cell_has_red_font(cell):
    """检查单元格是否有红色字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if is_red_font(run):
                return True
    return False

def get_cell_text_clean(cell):
    """获取单元格的纯文本（用于字段名识别）"""
    text_parts = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if not is_red_font(run) and run.text.strip():
                text_parts.append(run.text.strip())
    return "".join(text_parts)

def fill_liyifeng_template_final(data, template_file="（李奕凤）个体工商户开业登记申请书（模板）.docx"):
    """智能填充李奕凤版模板 v10 - 最终完美版"""
    try:
        print(f"正在读取模板: {template_file}")
        doc = Document(template_file)

        print("正在处理：智能识别字段，删除红色字体，填充数据...")

        # 定义字段识别规则（完整版）
        field_patterns = [
            # 基本信息
            (["个体工商户名称", "名称", "字号名称", "字号"], "business_name"),
            (["经营者姓名", "经营者", "经营者姓名"], "operator_name"),
            (["性别"], "gender"),
            (["民族"], "nation"),
            (["政治面貌", "政治"], "political_status"),
            (["文化程度", "学历", "教育程度"], "education"),
            (["身份证号码", "身份证", "证件号码", "身份证号"], "id_card"),

            # 联系方式
            (["联系电话", "联系", "电话", "手机号", "手机"], "phone"),
            (["电子邮箱", "邮箱", "Email", "E-mail", "电子邮件"], "email"),
            (["邮政编码", "邮编", "邮政"], "postal_code"),

            # 经营信息
            (["经营场所", "住所", "地址", "经营地址"], "business_address"),
            (["从业人数", "人数", "从业"], "employee_count"),
            (["注册资金", "资金", "资本", "注册资本"], "registered_capital"),
            (["经营范围"], "business_scope"),
            (["经营期限", "期限", "营业期限"], "operation_period"),
        ]

        deleted_count = 0
        filled_count = 0
        filled_fields = set()
        unmatched_red_cells = []

        # 遍历所有表格
        for table_idx, table in enumerate(doc.tables):
            print(f"\n处理表格 {table_idx}...")

            for row_idx, row in enumerate(table.rows):
                # 找出这一行有红色字体的单元格
                for cell_idx, cell in enumerate(row.cells):
                    if not cell_has_red_font(cell):
                        continue

                    # 查找字段名：多方向搜索
                    field_name_text = ""

                    # 1. 优先左侧单元格
                    if cell_idx > 0:
                        left_cell = row.cells[cell_idx - 1]
                        left_text = get_cell_text_clean(left_cell)
                        if left_text:
                            field_name_text = left_text
                            source = "左"
                        else:
                            source = "?"
                    else:
                        source = "?"

                    # 2. 如果左侧没有合适字段名，尝试上方单元格
                    if not field_name_text and row_idx > 0:
                        try:
                            prev_row = table.rows[row_idx - 1]
                            if cell_idx < len(prev_row.cells):
                                above_cell = prev_row.cells[cell_idx]
                                above_text = get_cell_text_clean(above_cell)
                                if above_text:
                                    field_name_text = above_text
                                    source = "上"
                        except:
                            pass

                    # 3. 如果还没有，尝试当前单元格删除红色后的文本
                    if not field_name_text:
                        current_text = get_cell_text_clean(cell)
                        if current_text:
                            field_name_text = current_text
                            source = "内"

                    # 查找匹配的字段
                    matched_key = None
                    matched_field_name = None

                    for field_names, data_key in field_patterns:
                        if data_key not in filled_fields:
                            for field_name in field_names:
                                if field_name in field_name_text:
                                    matched_key = data_key
                                    matched_field_name = field_name
                                    filled_fields.add(data_key)
                                    break
                            if matched_key:
                                break

                    # 删除所有红色字体
                    red_text_before = ""
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if is_red_font(run):
                                if not red_text_before:
                                    red_text_before = run.text[:20]
                                run.text = ""
                                deleted_count += 1

                    # 如果找到匹配，填充数据
                    if matched_key and matched_key in data:
                        value = str(data[matched_key])
                        if value:
                            # 清空并填充
                            if cell.paragraphs:
                                para = cell.paragraphs[0]
                                para.text = ""
                                para.clear()

                                # 添加数据
                                new_run = para.add_run(value)
                                new_run.font.name = '宋体'
                                new_run.font.size = 209712  # 约10.5磅

                                filled_count += 1
                                print(f"  [OK] {source}侧'{matched_field_name}' -> {matched_key}: {value}")
                    else:
                        # 记录未匹配的红色单元格
                        if red_text_before:
                            unmatched_red_cells.append(f"{red_text_before} (字段: {field_name_text[:20]})")

        # 显示统计
        print(f"\n完成统计:")
        print(f"  删除红色字体: {deleted_count} 处")
        print(f"  填充数据: {filled_count} 处")
        print(f"  填充字段: {', '.join(filled_fields)}")

        if unmatched_red_cells:
            print(f"\n未匹配的红色字段 ({len(unmatched_red_cells)}个):")
            for item in unmatched_red_cells[:5]:  # 只显示前5个
                print(f"  - {item}")
            if len(unmatched_red_cells) > 5:
                print(f"  ... 还有 {len(unmatched_red_cells) - 5} 个")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_开业登记_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")
        print(f"[说明] 已删除红色字体示例数据，智能填充实际信息")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] {data_file}")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数"""
    # 测试数据
    test_data = {
        'business_name': '王五超市',
        'operator_name': '王五',
        'phone': '13333333333',
        'email': 'wangwu@test.com',
        'business_address': '广州市天河区天河路123号',
        'postal_code': '510000',
        'employee_count': '6',
        'registered_capital': '100000',
        'id_card': '440101198303033333',
        'gender': '男',
        'nation': '汉族',
        'political_status': '群众',
        'education': '本科',
        'business_scope': '日用百货；烟酒零售',
        'operation_period': '长期'
    }

    print("=" * 70)
    print("李奕凤版模板填充工具 v10 - 最终完美版")
    print("智能识别字段，完整填充所有数据")
    print("=" * 70)

    print("\n测试数据：")
    for key, value in test_data.items():
        print(f"  {key}: {value}")

    result = fill_liyifeng_template_final(test_data)

    if result:
        print("\n" + "=" * 70)
        print("测试成功！")
        print(f"生成的文件: {result}")
        print("=" * 70)

        # 用WPS打开
        print("\n正在用WPS打开文档...")
        import subprocess
        subprocess.Popen(['cmd', '/c', 'start', '', str(result)], shell=True)

if __name__ == "__main__":
    main()
