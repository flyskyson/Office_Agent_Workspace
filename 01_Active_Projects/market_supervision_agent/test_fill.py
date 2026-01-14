#!/usr/bin/env python3
"""
测试填充功能 - 使用干净模板
"""

import sys
import json
import datetime
from pathlib import Path
from docx import Document

def safe_replace(text, old, new):
    """安全替换，处理编码问题"""
    try:
        # 清理文本中的代理字符
        clean_text = ''.join(char for char in text if ord(char) < 0xD800 or ord(char) > 0xDFFF)
        clean_old = ''.join(char for char in old if ord(char) < 0xD800 or ord(char) > 0xDFFF)
        clean_new = ''.join(char for char in new if ord(char) < 0xD800 or ord(char) > 0xDFFF)

        return clean_text.replace(clean_old, clean_new)
    except:
        return text

def fill_template_safe(data, template_file="个体工商户开业登记申请书_干净版.docx"):
    """安全填充模板"""
    try:
        print(f"正在读取模板: {template_file}")
        doc = Document(template_file)

        print("正在填充模板...")

        # 创建替换映射
        replacements = {
            "{{个体工商户名称}}": data.get("business_name", ""),
            "{{经营者姓名}}": data.get("operator_name", ""),
            "{{联系电话}}": data.get("phone", ""),
            "{{电子邮箱}}": data.get("email", ""),
            "{{经营场所}}": data.get("business_address", ""),
            "{{邮政编码}}": data.get("postal_code", ""),
            "{{从业人数}}": data.get("employee_count", ""),
            "{{注册资金}}": data.get("registered_capital", ""),
            "{{身份证号码}}": data.get("id_card", ""),
            "{{性别}}": data.get("gender", ""),
            "{{民族}}": data.get("nation", ""),
            "{{政治面貌}}": data.get("political_status", ""),
            "{{文化程度}}": data.get("education", ""),
            "{{经营范围}}": data.get("business_scope", ""),
            "{{经营期限}}": data.get("operation_period", ""),
            "{{签字日期}}": data.get("sign_date", ""),
            "{{申请日期}}": data.get("application_date", "")
        }

        # 在段落中替换
        replaced_count = 0
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            for old_text, new_text in replacements.items():
                if old_text in original_text:
                    paragraph.text = safe_replace(original_text, old_text, new_text)
                    replaced_count += 1
                    break

        # 在表格中替换（逐个单元格替换，保留格式）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    for old_text, new_text in replacements.items():
                        if old_text in original_text:
                            # 清理并替换
                            clean_new = ''.join(char for char in str(new_text) if ord(char) < 0xD800 or ord(char) > 0xDFFF)

                            # 保留单元格格式，只替换文本
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, clean_new)
                                        replaced_count += 1

        print(f"完成 {replaced_count} 处替换")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_申请书_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] {output_file}")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] {data_file}")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数"""
    # 测试数据
    test_data = {
        "business_name": "李四便利店",
        "operator_name": "李四",
        "phone": "13900139000",
        "email": "lisi@example.com",
        "business_address": "上海市浦东新区陆家嘴环路1000号",
        "postal_code": "200120",
        "employee_count": "5",
        "registered_capital": "100000",
        "id_card": "310101199002022345",
        "gender": "男",
        "nation": "汉族",
        "political_status": "群众",
        "education": "大学",
        "business_scope": "便利店经营；日用品销售",
        "operation_period": "长期",
        "application_date": datetime.datetime.now().strftime("%Y年%m月%d日"),
        "sign_date": datetime.datetime.now().strftime("%Y年%m月%d日")
    }

    print("=" * 60)
    print("测试填充功能")
    print("=" * 60)

    result = fill_template_safe(test_data)

    if result:
        print("\n" + "=" * 60)
        print("测试成功！")
        print(f"生成的文件: {result}")
        print("=" * 60)
    else:
        print("\n测试失败！")

if __name__ == "__main__":
    main()
