#!/usr/bin/env python3
"""
个体工商户开业登记申请书自动填充工具
版本: 1.0
日期: 2026-01-11
"""

import sys
from pathlib import Path
import json
import datetime

try:
    from docx import Document
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)

def check_template():
    """检查模板文件"""
    template_file = "个体工商户开业登记申请书（模板）.docx"
    if not Path(template_file).exists():
        print(f"错误: 模板文件不存在")
        print(f"请确保 '{template_file}' 在当前目录")
        return False
    return True

def collect_data():
    """收集用户数据"""
    print("\n请填写个体工商户信息:")
    print("=" * 50)

    data = {}

    # 必填字段
    required_fields = [
        ("个体工商户名称", "business_name"),
        ("经营者姓名", "operator_name"),
        ("身份证号码", "id_card"),
        ("经营场所", "business_address"),
        ("经营范围", "business_scope"),
        ("资金数额（元）", "registered_capital"),
        ("联系电话", "phone")
    ]

    for display_name, field_name in required_fields:
        while True:
            value = input(f"{display_name}: ").strip()
            if value:
                data[field_name] = value
                break
            print(f"  {display_name}是必填项")

    # 可选字段
    optional_fields = [
        ("电子邮箱", "email"),
        ("邮政编码", "postal_code"),
        ("经营面积（平方米）", "business_area"),
        ("主营业务", "main_business"),
        ("行业类型", "business_type"),
        ("资金来源", "capital_source"),
        ("出资形式", "investment_form"),
        ("经营期限", "operation_period")
    ]

    for display_name, field_name in optional_fields:
        value = input(f"{display_name}（按Enter跳过）: ").strip()
        if value:
            data[field_name] = value

    # 自动添加日期
    data["application_date"] = datetime.datetime.now().strftime("%Y年%m月%d日")
    data["fill_date"] = datetime.datetime.now().strftime("%Y年%m月%d日")

    return data

def fill_template(data):
    """填充模板"""
    try:
        # 读取模板
        template_file = "个体工商户开业登记申请书（模板）.docx"
        doc = Document(template_file)

        print("\n正在填充模板...")

        # 创建替换映射
        replacements = {
            # 基本字段
            "个体工商户名称": data.get("business_name", ""),
            "经营者姓名": data.get("operator_name", ""),
            "身份证号码": data.get("id_card", ""),
            "经营场所": data.get("business_address", ""),
            "经营范围": data.get("business_scope", ""),
            "资金数额": data.get("registered_capital", ""),
            "联系电话": data.get("phone", ""),

            # 可选字段
            "电子邮箱": data.get("email", ""),
            "邮政编码": data.get("postal_code", ""),
            "经营面积": data.get("business_area", ""),
            "主营业务": data.get("main_business", ""),
            "行业类型": data.get("business_type", ""),
            "资金来源": data.get("capital_source", ""),
            "出资形式": data.get("investment_form", ""),
            "经营期限": data.get("operation_period", ""),

            # 日期字段
            "申请日期": data.get("application_date", ""),
            "填写日期": data.get("fill_date", "")
        }

        # 在段落中替换
        replaced_count = 0
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text and new_text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
                    replaced_count += 1

        # 在表格中替换
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old_text, new_text in replacements.items():
                        if old_text in cell.text and new_text:
                            cell.text = cell.text.replace(old_text, new_text)
                            replaced_count += 1

        print(f"完成 {replaced_count} 处替换")

        # 生成输出文件名
        business_name = data.get("business_name", "未知")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("filled_applications")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / f"{business_name}_申请书_{timestamp}.docx"

        # 保存文件
        doc.save(str(output_file))

        print(f"\n[成功] 申请书生成成功！")
        print(f"[文件] 位置: {output_file}")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] 文件: {data_file}")

        return output_file

    except Exception as e:
        print(f"\n[错误] 填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def load_from_json():
    """从JSON文件加载数据"""
    json_file = input("请输入JSON文件路径: ").strip()
    if not Path(json_file).exists():
        print(f"文件不存在: {json_file}")
        return None

    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        print(f"从 {json_file} 加载数据成功")
        return data
    except Exception as e:
        print(f"加载失败: {str(e)}")
        return None

def main():
    """主函数"""
    print("=" * 60)
    print("个体工商户开业登记申请书自动填充工具")
    print("=" * 60)

    # 检查模板
    if not check_template():
        return

    print("\n请选择操作:")
    print("1. 交互式填写")
    print("2. 从JSON文件加载")
    print("3. 退出")

    choice = input("\n请输入选项 (1-3): ").strip()

    if choice == "1":
        data = collect_data()
    elif choice == "2":
        data = load_from_json()
        if not data:
            return
    elif choice == "3":
        print("\n再见！")
        return
    else:
        print("\n无效选项")
        return

    # 填充模板
    result = fill_template(data)

    if result:
        print("\n" + "=" * 50)
        print("完成！您可以在 filled_applications/ 目录找到生成的文件")
        print("=" * 50)

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n程序已取消")
    except Exception as e:
        print(f"\n程序错误: {str(e)}")
