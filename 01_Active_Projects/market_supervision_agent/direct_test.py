#!/usr/bin/env python3
"""
直接测试新版填充工具
"""

import sys
from pathlib import Path
import json
import datetime

try:
    from docx import Document
except ImportError:
    print("需要安装 python-docx 库")
    sys.exit(1)

def direct_test():
    """直接测试填充功能"""

    print("直接测试Word模板填充")
    print("=" * 60)

    # 检查文件
    template_file = "个体工商户开业登记申请书_新版.docx"
    test_data_file = "test_data.json"

    if not Path(template_file).exists():
        print(f"错误: 模板文件不存在")
        return False

    # 加载测试数据
    with open(test_data_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    print(f"测试数据: {data['business_name']}")

    try:
        # 读取模板
        print("\n1. 读取模板...")
        doc = Document(template_file)

        # 创建替换映射
        replacements = {
            "{{个体工商户名称}}": data.get("business_name", ""),
            "{{经营者姓名}}": data.get("operator_name", ""),
            "{{联系电话}}": data.get("phone", ""),
            "{{电子邮箱}}": data.get("email", ""),
            "{{经营场所}}": data.get("business_address", ""),
            "{{邮政编码}}": data.get("postal_code", ""),
            "{{从业人数}}": data.get("employee_count", ""),
            "{{注册资金}}": data.get("registered_capital", ""),
            "{{身份证号码}}": data.get("id_card", ""),
            "{{性别}}": data.get("gender", ""),
            "{{民族}}": data.get("nation", ""),
            "{{政治面貌}}": data.get("political_status", ""),
            "{{文化程度}}": data.get("education", ""),
            "{{经营范围}}": data.get("business_scope", ""),
            "{{经营期限}}": data.get("operation_period", ""),
            "{{签字日期}}": data.get("sign_date", ""),
            "{{申请日期}}": data.get("application_date", "")
        }

        # 在段落中替换
        print("2. 替换段落...")
        p_count = 0
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
                    p_count += 1

        # 在表格中替换
        print("3. 替换表格...")
        t_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old_text, new_text in replacements.items():
                        if old_text in cell.text:
                            cell.text = cell.text.replace(old_text, new_text)
                            t_count += 1

        print(f"   完成 {p_count + t_count} 处替换")

        # 保存文件
        print("4. 保存文件...")
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)

        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = output_dir / f"{data['business_name']}_测试_{timestamp}.docx"

        doc.save(str(output_file))

        print(f"\n[成功] 文件已生成: {output_file.name}")
        print(f"[位置] {output_file}")
        print(f"[大小] {output_file.stat().st_size / 1024:.1f} KB")

        # 保存数据
        data_file = output_file.with_suffix('.json')
        with open(data_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"[数据] {data_file.name}")

        return True

    except Exception as e:
        print(f"\n[错误] {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = direct_test()

    print("\n" + "=" * 60)
    if success:
        print("[成功] 测试通过！")
        print("\n下一步:")
        print("1. 打开 output/ 目录查看生成的Word文件")
        print("2. 检查内容是否正确填充")
        print("3. 可以使用自己的数据重新生成")
    else:
        print("[失败] 测试未通过")